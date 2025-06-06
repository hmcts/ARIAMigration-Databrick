from pyspark.sql import SparkSession
from pyspark.sql.functions import col, to_date, coalesce, greatest, lit, explode, count, when, current_date, input_file_name
from pyspark.sql.types import StructType, StructField, StringType, IntegerType, DateType, ArrayType, BooleanType
from docx import Document
from datetime import datetime
import re
import logging
import os
from typing import Dict, List, Any

# Create SparkSession
spark = SparkSession.builder \
    .appName("AppealsGoldOutputTesting") \
    .getOrCreate()

# Mount point for gold files
gold_mnt = "/mnt/gold/ARIADM/ARM/APPEALS"

# Initialize test variables
test_results = {
        "file_presence": {"status": True, "errors": [], "metrics": {}},
        "schema_validation": {"status": True, "errors": [], "metrics": {}},
        "json_content_validation": {"status": True, "errors": [], "metrics": {}},
        "a360_content_validation": {"status": True, "errors": [], "metrics": {}},
        "html_content_validation": {"status": True, "errors": [], "metrics": {}},
        "metadata_compliance": {"status": True, "errors": [], "metrics": {}},
        "cross_file_validation": {"status": True, "errors": [], "metrics": {}},
        "silver_transformation": {"status": True, "errors": [], "metrics": {}},
        "data_quality_metrics": {"status": True, "errors": [], "metrics": {}},
        "manifest_validation": {"status": True, "errors": [], "metrics": {}},
    "file_name_uniqueness": {"status": True, "errors": [], "metrics": {}},
    "metadata_validation": {"status": True, "errors": [], "metrics": {}},
    }

# Define gold output schemas
json_schema = StructType([
    StructField("CaseNo", StringType(), True),
    StructField("CasePrefix", StringType(), True),
    StructField("CaseYear", IntegerType(), True),
    StructField("CaseType", IntegerType(), True),
    StructField("AppealTypeId", IntegerType(), True),
    StructField("DateLodged", StringType(), True),
    StructField("PortId", IntegerType(), True),
    StructField("NationalityId", IntegerType(), True),
    StructField("Interpreter", StringType(), True),
    StructField("CountryId", IntegerType(), True),
    StructField("DateOfIssue", StringType(), True),
    StructField("FamilyCase", StringType(), True),
    StructField("OakingtonCase", StringType(), True),
    StructField("VisitVisaType", StringType(), True),
    StructField("AdditionalGrounds", StringType(), True),
    StructField("AppealCategories", StringType(), True),
    StructField("ThirdCountryId", IntegerType(), True),
    StructField("PubliclyFunded", StringType(), True),
    StructField("NonStandardSCPeriod", StringType(), True),
    StructField("CourtPreference", StringType(), True),
    StructField("ProvisionalDestructionDate", StringType(), True),
    StructField("FileInStatutoryClosure", StringType(), True),
    StructField("DateOfNextListedHearing", StringType(), True),
    StructField("DocumentsReceived", StringType(), True),
    StructField("OutOfTimeIssue", StringType(), True),
    StructField("ValidityIssues", StringType(), True),
    StructField("ReceivedFromRespondent", BooleanType(), True),
    StructField("DateAppealReceived", StringType(), True),
    StructField("InCamera", StringType(), True),
    StructField("DateOfApplicationDecision", StringType(), True),
    StructField("UserId", IntegerType(), True),
    StructField("HumanRights", StringType(), True),
    StructField("SecureCourtRequired", StringType(), True),
    StructField("POUShortName", StringType(), True),
    StructField("RespondentName", StringType(), True),
    StructField("LSCCommission", StringType(), True),
    StructField("AppellantId", IntegerType(), True),
    StructField("AppellantName", StringType(), True),
    StructField("AppellantForenames", StringType(), True),
    StructField("AppellantTitle", StringType(), True),
    StructField("AppellantBirthDate", StringType(), True),
    StructField("Detained", StringType(), True),
    StructField("DetentionCentre", StringType(), True),
    StructField("CentreTitle", StringType(), True),
    StructField("DetentionCentreType", IntegerType(), True),
    StructField("DCAddress1", StringType(), True),
    StructField("DCAddress2", StringType(), True),
    StructField("DCAddress3", StringType(), True),
    StructField("DCAddress4", StringType(), True),
    StructField("DCPostcode", StringType(), True),
    StructField("DCSdx", StringType(), True),
    StructField("HistoryDetails", ArrayType(StructType([
        StructField("HistoryId", IntegerType(), True),
        StructField("CaseNo", StringType(), True),
        StructField("HistDate", StringType(), True),
        StructField("HistType", IntegerType(), True),
        StructField("HistoryComment", StringType(), True),
        StructField("UserName", StringType(), True),
        StructField("UserType", StringType(), True),
        StructField("Fullname", StringType(), True),
        StructField("Extension", StringType(), True),
        StructField("DoNotUse", BooleanType(), True),
        StructField("HistTypeDescription", StringType(), True)
    ])), True),
    StructField("AppealCategoryDetails", ArrayType(StructType([
        StructField("CategoryDescription", StringType(), True),
        StructField("Flag", StringType(), True)
    ])), True),
    StructField("ASFReferenceNoStatus", StringType(), True),
    StructField("LCPRequested", StringType(), True),
    StructField("S17ReferenceStatus", StringType(), True),
    StructField("S20ReferenceStatus", StringType(), True),
    StructField("HomeOfficeWaiverStatus", StringType(), True),
    StructField("EmbassyLocation", StringType(), True),
    StructField("Embassy", StringType(), True),
    StructField("Surname", StringType(), True),
    StructField("Forename", StringType(), True),
    StructField("Title", StringType(), True),
    StructField("OfficialTitle", StringType(), True),
    StructField("EmbassyAddress1", StringType(), True),
    StructField("EmbassyAddress2", StringType(), True),
    StructField("DoNotUseEmbassy", BooleanType(), True),
    StructField("DedicatedHearingCentre", StringType(), True),
    StructField("Prefix", StringType(), True),
    StructField("CourtType", IntegerType(), True),
    StructField("HearingCentreAddress1", StringType(), True),
    StructField("HearingCentreAddress2", StringType(), True),
    StructField("HearingCentreAddress3", StringType(), True),
    StructField("HearingCentrePostcode", StringType(), True),
    StructField("HearingCentreTelephone", StringType(), True),
    StructField("HearingCentreFax", StringType(), True),
    StructField("HearingCentreSdx", StringType(), True),
    StructField("PouId", IntegerType(), True),
    StructField("MainLondonCentre", BooleanType(), True),
    StructField("DoNotUse", BooleanType(), True),
    StructField("CentreLocation", IntegerType(), True),
    StructField("OrganisationId", IntegerType(), True),
    StructField("Authorised", StringType(), True),
    StructField("AppealTypeDescription", StringType(), True),
    StructField("AppealTypePrefix", StringType(), True),
    StructField("AppealTypeNumber", IntegerType(), True),
    StructField("AppealTypeFullName", StringType(), True),
    StructField("AppealTypeCategory", IntegerType(), True),
    StructField("AppealType", IntegerType(), True),
    StructField("AppealTypeDoNotUse", BooleanType(), True),
    StructField("AppealTypeDateStart", StringType(), True),
    StructField("AppealTypeCategorieDetails", ArrayType(StructType([
        StructField("AppealTypeCategoryId", IntegerType(), True),
        StructField("AppealTypeId", IntegerType(), True),
        StructField("CategoryId", IntegerType(), True),
        StructField("FeeExempt", BooleanType(), True)
    ])), True),
    StructField("AppealGroundsDetails", ArrayType(StructType([
        StructField("AppealTypeId", IntegerType(), True),
        StructField("AppealTypeDescription", StringType(), True)
    ])), True)
])

a360_schema = StructType([
    StructField("operation", StringType(), True),
    StructField("relation_id", StringType(), True),
    StructField("record_metadata", StructType([
        StructField("publisher", StringType(), True),
        StructField("record_class", StringType(), True),
        StructField("region", StringType(), True),
        StructField("recordDate", StringType(), True),
        StructField("event_date", StringType(), True),
        StructField("client_identifier", StringType(), True),
        StructField("entitlement_tag", StringType(), True),
        StructField("bf_001", StringType(), True),  # HO Reference
        StructField("bf_002", StringType(), True),  # Forenames
        StructField("bf_003", StringType(), True),  # Surname
        StructField("bf_004", StringType(), True),  # Date of Birth
        StructField("bf_005", StringType(), True),  # Port Reference
        StructField("bf_006", StringType(), True)   # Postcode
    ]), True),
    StructField("file_metadata", StructType([
        StructField("publisher", StringType(), True),
        StructField("dz_file_name", StringType(), True),
        StructField("file_tag", StringType(), True)
    ]), True)
])

# Define Test suites

def test_file_presence():
    """Test presence of output files"""
    try:
        json_count = spark.read.format("binaryFile").load(f"{gold_mnt}/JSON").count()
        a360_count = spark.read.format("binaryFile").load(f"{gold_mnt}/A360").count()
        html_count = spark.read.format("binaryFile").load(f"{gold_mnt}/HTML").count()
        
        test_results["file_presence"]["metrics"].update({
            "json_files": json_count,
            "a360_files": a360_count,
            "html_files": html_count
        })
        
        if json_count == 0:
            test_results["file_presence"]["status"] = False
            test_results["file_presence"]["errors"].append("No JSON files found")
            
    except Exception as e:
        test_results["file_presence"]["status"] = False
        test_results["file_presence"]["errors"].append(str(e))

def test_schema_validation():
    """Validate schema of output files"""
    try:
        df_json = spark.read.json(f"{gold_mnt}/JSON/appeals_*.json")
        df_a360 = spark.read.json(f"{gold_mnt}/A360/appeals_*.a360")

        # Function to compare schemas and find missing fields
        def find_missing_fields(expected_schema, actual_schema, schema_name):
            expected_fields = {field.name for field in expected_schema}
            actual_fields = {field.name for field in actual_schema}
            missing_fields = expected_fields - actual_fields
            if missing_fields:
                return f"{schema_name} schema is missing fields: {missing_fields}"
            return None

        json_schema_error = find_missing_fields(json_schema, df_json.schema, "JSON")
        a360_schema_error = find_missing_fields(a360_schema, df_a360.schema, "A360")

        # Report schema mismatches
        if json_schema_error or a360_schema_error:
            test_results["schema_validation"]["status"] = False
            if json_schema_error:
                test_results["schema_validation"]["errors"].append(json_schema_error)
            if a360_schema_error:
                test_results["schema_validation"]["errors"].append(a360_schema_error)

    except Exception as e:
        test_results["schema_validation"]["status"] = False
        test_results["schema_validation"]["errors"].append(str(e))

# def test_schema_validation():
#     """Validate schema of output files"""
#     try:
#         df_json = spark.read.json(f"{gold_mnt}/JSON/appeals_*.json")
#         df_a360 = spark.read.json(f"{gold_mnt}/A360/appeals_*.a360")
        
#         # Compare schemas
#         if df_json.schema != json_schema:
#             test_results["schema_validation"]["status"] = False
#             test_results["schema_validation"]["errors"].append("JSON schema mismatch")
            
#         if df_a360.schema != a360_schema:
#             test_results["schema_validation"]["status"] = False
#             test_results["schema_validation"]["errors"].append("A360 schema mismatch")
            
#     except Exception as e:
#         test_results["schema_validation"]["status"] = False
#         test_results["schema_validation"]["errors"].append(str(e))

def test_json_content():
    """
    Validates the content of JSON files against expected business rules and data quality standards.
    
    Tests:
    - Mandatory field presence
    - Date format validity
    - Reference data integrity
    - Business logic rules
    - Data relationships
    - Data range validation
    """
    try:
        df_json = spark.read.json(f"{gold_mnt}/JSON/appeals_*.json", schema=json_schema)

        validation_results = {
            "mandatory_fields": {},
            "date_validations": {},
            "reference_integrity": {},
            "business_rules": {},
            "data_relationships": {},
            "data_range": {},
            "consistency": {}
        }

        test_results["json_content_validation"] = {
            "status": True,
            "errors": [],
            "metrics": {}
        }

        # Mandatory field validation
        mandatory_fields = [
            "CaseNo", "CasePrefix", "CaseYear", "AppealTypeId", "DateLodged",
            "AppellantName", "AppellantForenames", "AppellantBirthDate"
        ]

        for field in mandatory_fields:
            null_count = df_json.filter(col(field).isNull()).count()
            validation_results["mandatory_fields"][field] = null_count

            if null_count > 0:
                test_results["json_content_validation"]["status"] = False
                test_results["json_content_validation"]["errors"].append(
                    f"Mandatory field '{field}' has {null_count} null values"
                )

        validation_results["mandatory_fields"].update({
            "mandatory_field_errors": null_count
        })

        # Date format validation
        date_fields = [
            "DateLodged", "DateOfIssue", "ProvisionalDestructionDate",
            "DateAppealReceived", "DateOfApplicationDecision", "AppellantBirthDate"
        ]

        for field in date_fields:
            invalid_dates = df_json.filter(
                ~to_date(col(field), "yyyy-MM-dd'T'HH:mm:ss.SSS'Z'").isNull() &
                col(field).isNotNull()
            ).count()

            validation_results["date_validations"][field] = invalid_dates

            if invalid_dates > 0:
                test_results["json_content_validation"]["status"] = False
                test_results["json_content_validation"]["errors"].append(
                    f"Date field '{field}' has {invalid_dates} invalid format entries"
                )

        validation_results["date_validations"].update({
            "date_validation_errors": invalid_dates
        })

        # Reference data validation
        ref_conditions = [
            (col("CasePrefix").isin(["VA", "AA", "AS", "CC", "HR", "HX", "IM", "NS", "OA", "OC", "RD", "TH", "XX", "DA", "DC", "EA", "HU", "PA", "RP", "LP", "LR", "LD", "LH", "LE", "IA"])),
            (col("AppealTypeId").isNotNull()),
            (col("Interpreter").isin(["YES", "NO", "disabled"])),
            (col("HumanRights").isin(["YES", "NO", "disabled"]))
        ]

        for condition in ref_conditions:
            invalid_count = df_json.filter(~condition).count()
            if invalid_count > 0:
                test_results["json_content_validation"]["status"] = False
                test_results["json_content_validation"]["errors"].append(
                    f"Reference data validation failed with {invalid_count} invalid records"
                )

        validation_results["reference_integrity"].update({
            "reference_integrity_errors": invalid_count
        })

        # Business logic validation
        # business_rules = [
        #     (col("CaseYear").between(1900, 2025)),
        #     (col("DateLodged") <= current_date()),
        #     (when(col("Interpreter") == "YES", col("CourtPreference").isNotNull()).otherwise(lit(True))),
        #     (when(col("CaseStatus").isin(["40", "41", "42", "43", "44", "45", "53", "27", "28", "29", "34", "32", "33"]) & (col("Outcome") == 86), lit(True)).otherwise(lit(False))),
        #     (when(col("CaseStatus").isin(["40", "41", "42", "43", "44", "45", "53", "27", "28", "29", "34", "32", "33"]) & (col("Outcome") == 0), lit(True)).otherwise(lit(False))),
        # ]

        # for rule in business_rules:
        #     violations = df_json.filter(~rule).count()
        #     if violations > 0:
        #         test_results["json_content_validation"]["status"] = False
        #         test_results["json_content_validation"]["errors"].append(
        #             f"Business rule violation found in {violations} records"
        #         )

        # validation_results["business_rules"].update({
        #     "business_rules_errors": violations
        # })

        # Data relationship validation
        relationship_conditions = [
            (col("AppellantId").isNotNull()),
            (col("AppealTypeId").isNotNull())
        ]

        for condition in relationship_conditions:
            invalid_count = df_json.filter(~condition).count()
            if invalid_count > 0:
                test_results["json_content_validation"]["status"] = False
                test_results["json_content_validation"]["errors"].append(
                    f"Data relationship validation failed with {invalid_count} invalid records"
                )

        validation_results["data_relationships"].update({
            "data_relationships_errors": invalid_count
        })

        # Data range validation
        # range_conditions = [
        #     (col("CaseYear").between(1900, 2025)),
        #     (col("DetentionCentreType").isin([1, 2, 3])),
        #     (col("CaseStatus").isin(["40", "41", "42", "43", "44", "45", "53", "27", "28", "29", "34", "32", "33"])),
        #     (col("Outcome").isin(["0", "86", "109", "104", "82", "99", "121", "27", "39", "1", "50", "40", "52", "89", "37", "5", "30", "31", "25", "14", "80", "94", "93", "91", "95", "108"])),
        # ]

        # for condition in range_conditions:
        #     invalid_count = df_json.filter(~condition).count()
        #     if invalid_count > 0:
        #         test_results["json_content_validation"]["status"] = False
        #         test_results["json_content_validation"]["errors"].append(
        #             f"Data range validation failed with {invalid_count} invalid records"
        #         )

        # validation_results["data_ranges"].update({
        #     "data_ranges_errors": invalid_count
        # })

        # Consistency with segmentation validation
        # consistency_conditions = [
        #     (when((col("CasePrefix").isin(["DA", "DC", "EA", "HU", "PA", "RP"]) | 
        #         (col("CasePrefix").isin(["LP", "LR", "LD", "LH", "LE", "IA"]) & col("HOANRef").isNull())) &
        #         (col("CaseStatus").isNull() | 
        #         (col("CaseStatus") == "10") & (col("Outcome").isin(["0", "109", "104", "82", "99", "121", "27", "39"])) |
        #         (col("CaseStatus") == "46") & (col("Outcome").isin(["1", "86"])) |
        #         (col("CaseStatus") == "26") & (col("Outcome").isin(["0", "27", "39", "50", "40", "52", "89"])) |
        #         (col("CaseStatus").isin(["37", "38"])) & (col("Outcome").isin(["39", "40", "37", "50", "27", "0", "5", "52"])) |
        #         (col("CaseStatus") == "39") & (col("Outcome").isin(["0", "86"])) |
        #         (col("CaseStatus") == "50") & (col("Outcome") == "0") |
        #         (col("CaseStatus").isin(["52", "36"])) & (col("Outcome") == "0") & (col("st.DecisionDate").isNull())), lit(True)).otherwise(lit(False)),
        # )]

        # for condition in consistency_conditions:
        #     invalid_count = df_json.filter(~condition).count()
        #     if invalid_count > 0:
        #         test_results["json_content_validation"]["status"] = False
        #         test_results["json_content_validation"]["errors"].append(
        #             f"Consistency validation failed with {invalid_count} invalid records"
        #         )

        # validation_results["consistency"].update({
        #     "consistency_errors": invalid_count
        # })

        # Update test results
        for category, results in validation_results.items():
            has_failures = any(v > 0 for v in results.values() if isinstance(v, (int, float)))
            if has_failures:
                test_results["json_content_validation"]["status"] = False
                test_results["json_content_validation"]["errors"].append(
                    f"Validation failures found in {category}: {results}"
                )

        # Update metrics
        test_results["json_content_validation"]["metrics"].update(validation_results)

    except Exception as e:
        test_results["json_content_validation"] = {
            "status": False,
            "errors": [f"JSON content validation failed: {str(e)}"],
            "metrics": {}
        }

def test_a360_content():
    """
    Validates the content and structure of A360 files against expected business rules and data quality standards.
    
    Tests:
    - Record creation operations
    - File upload operations 
    - Metadata compliance
    - Field formatting
    - Reference data integrity
    - File associations
    """
    try:
        df_a360 = spark.read.json(f"{gold_mnt}/A360/appeals_*.a360")
        
        a360_validation_results = {
            "record_operations": {},
            "file_operations": {},
            "metadata_compliance": {},
            "field_formatting": {},
            "file_associations": {}
        }

        # Record Creation Operation Validation
        create_records = df_a360.filter(col("operation") == "create_record")
        upload_files = df_a360.filter(col("operation") == "upload_new_file")

        # Check required fields for create_record operations
        create_record_validation = create_records.select([
            count(when(col("relation_id").isNull(), True)).alias("missing_relation_id"),
            count(when(col("record_metadata.publisher").isNull(), True)).alias("missing_publisher"),
            count(when(col("record_metadata.record_class").isNull(), True)).alias("missing_record_class"),
            count(when(col("record_metadata.region").isNull(), True)).alias("missing_region"),
            count(when(col("record_metadata.recordDate").isNull(), True)).alias("missing_record_date"),
            count(when(col("record_metadata.event_date").isNull(), True)).alias("missing_event_date"),
            count(when(col("record_metadata.client_identifier").isNull(), True)).alias("missing_client_id"),
            count(when(col("record_metadata.entitlement_tag").isNull(), True)).alias("missing_entitlement_tag")
        ]).collect()[0]

        a360_validation_results["record_operations"].update({
            "total_records": create_records.count(),
            "validation_failures": create_record_validation
        })

        # Metadata Compliance Checks
        metadata_validation = create_records.select([
            count(when(col("record_metadata.publisher") != "ARIA", True)).alias("invalid_publisher"),
            count(when(col("record_metadata.record_class").isin(["ARIAFTA", "ARIAUTA"]), True)).alias("invalid_record_class"),
            count(when(col("record_metadata.region") != "GBR", True)).alias("invalid_region"),
            count(when(~col("record_metadata.client_identifier").rlike("^[A-Z]{2}/\\d{5}/\\d{4}$"), True)).alias("invalid_client_id_format"),
            count(when(col("record_metadata.entitlement_tag") != "IA_Tribunal", True)).alias("invalid_entitlement_tag"),
            count(when(col("record_metadata.event_date").isNull(), True)).alias("missing_event_date"),
            count(when(col("record_metadata.recordDate").isNull(), True)).alias("missing_record_date"),
        ]).collect()[0]

        a360_validation_results["metadata_compliance"].update({
            "validation_failures": metadata_validation
        })

        # File Upload Operation Validation
        file_upload_validation = upload_files.select([
            count(when(col("relation_id").isNull(), True)).alias("missing_relation_id"),
            count(when(col("file_metadata.publisher").isNull(), True)).alias("missing_publisher"),
            count(when(col("file_metadata.dz_file_name").isNull(), True)).alias("missing_filename"),
            count(when(col("file_metadata.file_tag").isNull(), True)).alias("missing_file_tag")
        ]).collect()[0]

        a360_validation_results["file_operations"].update({
            "total_files": upload_files.count(),
            "validation_failures": file_upload_validation
        })

        # File Tag Validation
        file_tag_validation = upload_files.select([
            count(when(~col("file_metadata.file_tag").isin(["html", "json"]), True)).alias("invalid_file_tag")
        ]).collect()[0]

        a360_validation_results["file_operations"].update({
            "file_tag_failures": file_tag_validation
        })

        # Business Field Validation
        # business_field_validation = create_records.select([
        #     count(when(col("record_metadata.bf_001").isNull() & 
        #               col("record_metadata.bf_002").isNull() & 
        #               col("record_metadata.bf_003").isNull(), True)).alias("missing_all_bf_fields"),
        #     count(when(~to_date(col("record_metadata.bf_004"), "yyyy-MM-dd HH:mm:ss").isNull() & 
        #               col("record_metadata.bf_004") != "None", True)).alias("valid_date_format")
        # ]).collect()[0]

        # a360_validation_results["field_formatting"].update({
        #     "business_field_failures": business_field_validation
        # })

        # File Association Validation
        file_associations = (
            upload_files.groupBy("relation_id")
            .agg(
                count(when(col("file_metadata.file_tag") == "html", True)).alias("html_count"),
                count(when(col("file_metadata.file_tag") == "json", True)).alias("json_count")
            )
            .filter((col("html_count") != 1) | (col("json_count") != 1))
            .count()
        )

        a360_validation_results["file_associations"].update({
            "invalid_associations": file_associations
        })

        # Update test results
        for category, results in a360_validation_results.items():
            has_failures = any(v > 0 for v in results.values() if isinstance(v, (int, float)))
            if has_failures:
                test_results["a360_content_validation"]["status"] = False
                test_results["a360_content_validation"]["errors"].append(
                    f"Validation failures found in {category}: {results}"
                )
            else:
                test_results["a360_content_validation"]["status"] = True

        test_results["a360_content_validation"]["metrics"].update(a360_validation_results)

    except Exception as e:
        test_results["a360_content_validation"]["status"] = False
        test_results["a360_content_validation"]["errors"].append(f"A360 content validation failed: {str(e)}")

def test_html_content():
    """
    Validates the content and structure of HTML archive files against expected business rules and data quality standards.
    
    Tests:
    - Document structure
    - Required sections presence
    - Field availability
    - Data consistency
    - Special section handling
    - Document formatting
    """
    try:
        df_html = spark.read.text(f"{gold_mnt}/HTML/appeals_*.html").withColumn("file_path", input_file_name())
        
        validation_results = {
            "document_structure": {},
            "required_sections": {},
            "field_presence": {},
            "data_consistency": {},
            "special_sections": {},
            "formatting": {}
        }

        # Define expected sections and subsections
        expected_sections = [
            "Case details",
            "Main",
            "Parties",
            "Status",
            "Status Details",
            "Outcome",
            "Adjournment/Withdrawal Details",
            "Hearing Details",
            "Review Directions",
            "History",
            "B/F Diary",
            "S.C. / Misc",
            "Cost order details",
            "Linked Files",
            "Maintain Typing",
            "Additional Grounds",
            "Human Rights",
            "Appeal Categories",
            "New Matters",
            "Document Tracking",
            "Respondent",
            "Representative",
            "Sponsor",
            "Dependents",
            "Payments"
        ]

        # Define critical fields that must be present
        critical_fields = [
            "Case No :",
            "Name :",
            "Forenames :",
            "Date of Birth :",
            "HO Ref :",
            "Appeal type :",
            "Lodged :",
            "Received :",
            "Status of case :",
            "Date of hearing :",
            "Interpreter(s) required :"
        ]

        # Parse HTML content and perform validations
        for html_file in df_html.collect():
            content = html_file.value
            file_path = html_file.file_path
            file_name = file_path.split("/")[-1] # Extract filename

            # Document Structure Validation
            structure_validation = {
                "has_header": "ARIA Archived Case Record" in content,
                "has_footer": "Notes:" in content,
                "section_count": sum(1 for section in expected_sections if section in content)
            }
            
            validation_results["document_structure"].update(structure_validation)
            # validation_results["document_structure"][file_name] = structure_validation #store by file

            # Required Sections Validation
            missing_sections = [section for section in expected_sections if section not in content]
            validation_results["required_sections"] = {
                "missing_sections": missing_sections,
                "total_missing": len(missing_sections),
                "html_file": file_name
            }

            # Field Presence Validation
            missing_fields = [field for field in critical_fields if field.lower() not in content.lower()]
            validation_results["field_presence"].update({
                "missing_critical_fields": missing_fields,
                "total_missing": len(missing_fields),
                "html_file": file_name
            })

            # Special Sections Validation
            special_sections_validation = {
                "has_history": "History" in content and "Event type" in content,
                "has_bf_diary": "B/F Diary" in content and "B/F Date" in content,
                "has_payment_events": "Payment Events Summary" in content,
                "has_appeal_categories": "Appeal Categories" in content,
                "has_interpreter_requirements": "Interpreter Requirements" in content
            }
            
            validation_results["special_sections"].update(special_sections_validation)

            # Data Consistency Validation
            consistency_checks = {
                "valid_case_number": bool(re.search(r"Case No\s*:\s*[A-Z]{2}/\d{5}/\d{4}", content)),
                "valid_dates": all(
                    re.match(r"\d{2}/\d{2}/\d{4}", date) 
                    for date in re.findall(r"\d{2}/\d{2}/\d{4}", content)
                ),
                "matching_appellant_details": "Main Appellant" in content and "Name :" in content
            }
            
            validation_results["data_consistency"].update(consistency_checks)

            # Document Formatting Validation
            formatting_checks = {
                "proper_section_breaks": content.count("\n\n") > 10,
                "consistent_field_separators": content.count(" : ") > 20,
                "table_formatting": "Status No." in content and "Status" in content
            }
            
            validation_results["formatting"].update(formatting_checks)

        # Evaluate validation results
        for category, results in validation_results.items():
            if isinstance(results, dict):
                has_failures = any(
                    (isinstance(v, bool) and not v) or
                    (isinstance(v, (list, tuple)) and len(v) > 0) or
                    (isinstance(v, (int, float)) and v > 0)
                    for v in results.values()
                )
                
                if has_failures:
                    test_results["html_content_validation"]["status"] = False
                    test_results["html_content_validation"]["errors"].append(
                        f"Validation failures found in {category}: {results}"
                    )

        test_results["html_content_validation"]["metrics"].update(validation_results)

    except Exception as e:
        test_results["html_content_validation"]["status"] = False
        test_results["html_content_validation"]["errors"].append(f"HTML content validation failed: {str(e)}")


def test_manifest_files():
    """
    Validates the structure, content, and batch size of the manifest files stored in the gold layer.
    """
    try:
        manifest_files = spark.read.json(f"{gold_mnt}/MANIFEST/appeals_*.manifest")
        
        # Validate manifest file structure
        expected_manifest_schema = StructType([
            StructField("sequence_number", IntegerType(), True),
            StructField("case_id", StringType(), True),
            StructField("case_type", StringType(), True),
            StructField("file_name", StringType(), True),
            StructField("file_type", StringType(), True),
            StructField("file_size", IntegerType(), True),
            StructField("file_hash", StringType(), True),
            StructField("created_date", StringType(), True)
        ])
        
        if manifest_files.schema != expected_manifest_schema:
            test_results["manifest_validation"]["status"] = False
            test_results["manifest_validation"]["errors"].append("Manifest file schema mismatch")
        
        # Validate batch size
        batch_size = manifest_files.groupBy("sequence_number").count().collect()
        invalid_batches = [row["sequence_number"] for row in batch_size if row["count"] != 250]
        
        if invalid_batches:
            test_results["manifest_validation"]["status"] = False
            test_results["manifest_validation"]["errors"].append(f"Invalid batch sizes found for sequence numbers: {invalid_batches}")
        
        # Validate case_id format
        invalid_case_ids = manifest_files.filter(~col("case_id").rlike("^[A-Z]{2}/\\d{5}/\\d{4}$")).count()
        
        if invalid_case_ids > 0:
            test_results["manifest_validation"]["status"] = False
            test_results["manifest_validation"]["errors"].append(f"Found {invalid_case_ids} invalid case IDs in manifest files")
        
        test_results["manifest_validation"]["metrics"].update({
            "total_manifest_files": manifest_files.select("sequence_number").distinct().count(),
            "total_records": manifest_files.count()
        })
        
    except Exception as e:
        test_results["manifest_validation"]["status"] = False
        test_results["manifest_validation"]["errors"].append(f"Manifest validation failed: {str(e)}")


def test_file_name_uniqueness():
    """
    Validates the uniqueness of file names for HTML and JSON files stored in the gold layer.
    """
    try:
        html_files = spark.read.format("binaryFile").load(f"{gold_mnt}/HTML")
        json_files = spark.read.format("binaryFile").load(f"{gold_mnt}/JSON")
        
        # Check for duplicate file names in HTML files
        html_duplicates = html_files.groupBy("path").count().filter("count > 1").count()
        
        if html_duplicates > 0:
            test_results["file_name_uniqueness"]["status"] = False
            test_results["file_name_uniqueness"]["errors"].append(f"Found {html_duplicates} duplicate file names in HTML files")
        
        # Check for duplicate file names in JSON files
        json_duplicates = json_files.groupBy("path").count().filter("count > 1").count()
        
        if json_duplicates > 0:
            test_results["file_name_uniqueness"]["status"] = False
            test_results["file_name_uniqueness"]["errors"].append(f"Found {json_duplicates} duplicate file names in JSON files")
        
    except Exception as e:
        test_results["file_name_uniqueness"]["status"] = False
        test_results["file_name_uniqueness"]["errors"].append(f"File name uniqueness validation failed: {str(e)}")

def test_metadata_validation():
    """
    Validates the presence and accuracy of mandatory metadata fields in the generated JSON files.
    """
    try:
        df_json = spark.read.json(f"{gold_mnt}/JSON/appeals_*.json", schema=json_schema)
        
        # Validate mandatory metadata fields
        mandatory_metadata_fields = [
            "CaseNo",
            "DateOfIssue",
            "DateLodged",
            "region",
            "publisher",
            "record_class",
            "entitlement_tag"
        ]
        
        for field in mandatory_metadata_fields:
            missing_count = df_json.filter(col(field).isNull()).count()
            
            if missing_count > 0:
                test_results["metadata_validation"]["status"] = False
                test_results["metadata_validation"]["errors"].append(f"Mandatory metadata field '{field}' has {missing_count} missing values")
        
        # Validate metadata field formats and values
        metadata_validations = [
            (col("client_identifier").rlike("^[A-Z]{2}/\\d{5}/\\d{4}$"), "Invalid client_identifier format"),
            (col("event_date").cast("date").isNotNull(), "Invalid event_date format"),
            (col("recordDate").cast("date").isNotNull(), "Invalid recordDate format"),
            (col("region") == "GBR", "Invalid region value"),
            (col("publisher") == "ARIA", "Invalid publisher value"),
            (col("record_class").isin(["ARIAFTA", "ARIAUTA"]), "Invalid record_class value"),
            (col("entitlement_tag") == "IA_Tribunal", "Invalid entitlement_tag value")
        ]
        
        for condition, error_message in metadata_validations:
            invalid_count = df_json.filter(~condition).count()
            
            if invalid_count > 0:
                test_results["metadata_validation"]["status"] = False
                test_results["metadata_validation"]["errors"].append(f"{error_message}: {invalid_count} invalid records")
        
        # Validate business metadata fields
        business_metadata_fields = ["bf_001", "bf_002", "bf_003", "bf_004", "bf_005", "bf_006"]
        
        for field in business_metadata_fields:
            missing_count = df_json.filter(col(field).isNull()).count()
            
            if missing_count > 0:
                test_results["metadata_validation"]["status"] = False
                test_results["metadata_validation"]["errors"].append(f"Business metadata field '{field}' has {missing_count} missing values")
        
    except Exception as e:
        test_results["metadata_validation"]["status"] = False
        test_results["metadata_validation"]["errors"].append(f"Metadata validation failed: {str(e)}")

def run_all_tests():
    """Execute all tests"""
    # Execute test suites
    test_file_presence()
    test_schema_validation()
    test_json_content()
    test_a360_content()
    test_file_name_uniqueness()
    test_manifest_files()
    test_metadata_validation()
    test_html_content()
    # Generate comprehensive report
    generate_test_report(test_results)

def generate_test_report(test_results):
    """
    Generates a comprehensive test report in Microsoft Word format.
    
    Parameters:
        test_results (dict): Dictionary containing all test results and metrics
    
    The report includes:
    - Executive summary
    - Test execution details
    - Detailed results by category
    - Data quality metrics
    """
    document = Document()
    
    # Title and Header
    document.add_heading('Appeals Gold Output Testing Results', 0)
    document.add_paragraph(f'Test Execution Date: {datetime.now().strftime("%d %B %Y %H:%M:%S")}')
    
    # Executive Summary
    document.add_heading('Executive Summary', 1)
    all_passed = all(result["status"] for result in test_results.values())
    summary_status = "✓ All validation tests passed successfully" if all_passed else "✗ Some validation tests failed"
    document.add_paragraph(summary_status)
    
    # Test Execution Details
    document.add_heading('Test Execution Details', 1)
    document.add_paragraph('The following validation suites were executed:')
    for test_name in test_results.keys():
        document.add_paragraph(f'• {test_name.replace("_", " ").title()}', style='List Bullet')
    
    # Detailed Results
    document.add_heading('Detailed Results', 1)
    
    for test_name, result in test_results.items():
        section = document.add_heading(test_name.replace("_", " ").title(), 2)
        
        # Status Summary
        status = "✓ Passed" if result["status"] else "✗ Failed"
        document.add_paragraph(f"Status: {status}")
        
        # Metrics Section
        if result["metrics"]:
            metrics_heading = document.add_heading("Metrics", 3)
            
            # Handle nested metrics
            def add_metrics(metrics, level=0):
                for metric, value in metrics.items():
                    if isinstance(value, dict):
                        document.add_paragraph(f"{metric}:", style=f'Heading {min(level + 4, 9)}')
                        add_metrics(value, level + 1)
                    else:
                        metric_text = "  " * level + f"• {metric}: {value}"
                        document.add_paragraph(metric_text, style='List Bullet')
            
            add_metrics(result["metrics"])
        
        # Errors Section
        if result["errors"]:
            errors_heading = document.add_heading("Errors", 3)
            for error in result["errors"]:
                document.add_paragraph(f"• {error}", style='List Bullet')
    
    # Data Quality Metrics
    document.add_heading('Data Quality Metrics', 1)
    if "data_quality_metrics" in test_results:
        for metric, value in test_results["data_quality_metrics"]["metrics"].items():
            document.add_paragraph(f"• {metric}: {value}", style='List Bullet')
    
    # Recommendations
    document.add_heading('Recommendations', 1)
    if not all_passed:
        document.add_paragraph('Based on the test results, the following actions are recommended:')
        for test_name, result in test_results.items():
            if not result["status"]:
                document.add_paragraph(
                    f'• Review and address issues in {test_name.replace("_", " ").title()}',
                    style='List Bullet'
                )
    else:
        document.add_paragraph('All tests passed successfully. No immediate actions required.')
    
    # Save report with timestamp
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    report_filename = f'appeals_gold_output_test_results_{timestamp}.docx'
    document.save(report_filename)
    
    return report_filename


if __name__ == "__main__":
    run_all_tests()
