from functools import reduce
from pyspark.sql import SparkSession
from pyspark.sql.functions import col, to_date, coalesce, greatest, lit, explode, date_add, current_date, count
from pyspark.sql.types import (
    StructType,
    StructField,
    StringType,
    IntegerType,
    DateType,
    ArrayType,
)
from docx import Document
from docx.shared import Inches

# Create a SparkSession
spark = SparkSession.builder.appName("BailsGoldOutputTesting").getOrCreate()

# Mount point for the gold files
gold_mnt = "/mnt/gold/ARIADM/ARM/BAILS"
bronze_mnt = "/mnt/bronze/ARIADM/ARM/BAILS"

# Initialise test variables for M1-M8 checks
m1_test = True
m2_test = True
m3_test = True
m4_test = True
m5_test = True
m6_test = True
m7_test = True
m8_test = True

m1_diff = ""
m2_diff = ""
m3_diff = ""
m4_diff = ""
m5_diff = ""
m6_diff = ""
m7_diff = ""
m8_diff = ""

# Define expected columns for each M query
expected_m1_columns = [
    "CaseNo",
    "HORef",
    "BailType",
    "CourtPreference",
    "DateOfIssue",
    "DateReceived",
    "FileLocationNote",
    "Language",
    "DedicatedHearingCentre"
]

expected_m2_columns = [
    "CaseNo",
    "AppellantId",
    "PortReference",
    "AppellantName",
    "AppellantForenames",
    "AppellantBirthDate",
    "AppellantDetained"
]

expected_m3_columns = [
    "CaseNo",
    "StatusId",
    "HearingDate",
    "HearingTypeDesc",
    "HearingCentreDesc",
    "Chairman",
    "AdjudicatorSurname"
]

expected_m4_columns = [
    "CaseNo",
    "BFDate",
    "Entry",
    "DateCompleted",
    "BFTypeDescription"
]

expected_m5_columns = ["CaseNo", "HistoryId", "HistDate", "HistType", "HistoryComment"]

expected_m6_columns = ["LinkNo", "CaseNo", "Name", "Forenames", "LinkDetailComment"]

expected_m7_columns = [
    "StatusId",
    "CaseNo",
    "DecisionDate",
    "ResidenceOrder",
    "ReportingOrder",
    "BailConditions",
    "InterpreterRequired",
]

expected_m8_columns = [
    "CaseNo", 
    "CategoryDescription", 
    "Flag", 
    "Priority"
]

# Test M1-M8 transformations
try:
    # Read M1-M8 tables
    m1_path = f"{bronze_mnt}/bronze_bail_ac_cr_cs_ca_fl_cres_mr_res_lang"
    df_m1 = spark.read.format("delta").load(m1_path)
    actual_m1_columns = df_m1.columns
    if set(expected_m1_columns) - set(actual_m1_columns):
        m1_test = False
        m1_diff = f"Missing columns in M1 output: {set(expected_m1_columns) - set(actual_m1_columns)}"

    m2_path = f"{bronze_mnt}/bronze_bail_ac_ca_apt_country_detc"
    df_m2 = spark.read.format("delta").load(m2_path)
    actual_m2_columns = df_m2.columns
    if set(expected_m2_columns) - set(actual_m2_columns):
        m2_test = False
        m2_diff = f"Missing columns in M2 output: {set(expected_m2_columns) - set(actual_m2_columns)}"

    m3_path = f"{bronze_mnt}/bronze_bail_ac_cl_ht_list_lt_hc_c_ls_adj"
    df_m3 = spark.read.format("delta").load(m3_path)
    actual_m3_columns = df_m3.columns
    if set(expected_m3_columns) - set(actual_m3_columns):
        m3_test = False
        m3_diff = f"Missing columns in M3 output: {set(expected_m3_columns) - set(actual_m3_columns)}"

    m4_path = f"{bronze_mnt}/bronze_bail_ac_bfdiary_bftype"
    df_m4 = spark.read.format("delta").load(m4_path)
    actual_m4_columns = df_m4.columns
    if set(expected_m4_columns) - set(actual_m4_columns):
        m4_test = False
        m4_diff = f"Missing columns in M4 output: {set(expected_m4_columns) - set(actual_m4_columns)}"

    m5_path = f"{bronze_mnt}/bronze_bail_ac_history_users"
    df_m5 = spark.read.format("delta").load(m5_path)
    actual_m5_columns = df_m5.columns
    if set(expected_m5_columns) - set(actual_m5_columns):
        m5_test = False
        m5_diff = f"Missing columns in M5 output: {set(expected_m5_columns) - set(actual_m5_columns)}"

    m6_path = f"{bronze_mnt}/bronze_bail_ac_link_linkdetail"
    df_m6 = spark.read.format("delta").load(m6_path)
    actual_m6_columns = df_m6.columns
    if set(expected_m6_columns) - set(actual_m6_columns):
        m6_test = False
        m6_diff = f"Missing columns in M6 output: {set(expected_m6_columns) - set(actual_m6_columns)}"

    m7_path = f"{bronze_mnt}/bronze_bail_status_sc_ra_cs"
    df_m7 = spark.read.format("delta").load(m7_path)
    actual_m7_columns = df_m7.columns
    if set(expected_m7_columns) - set(actual_m7_columns):
        m7_test = False
        m7_diff = f"Missing columns in M7 output: {set(expected_m7_columns) - set(actual_m7_columns)}"

    m8_path = f"{bronze_mnt}/bronze_bail_ac_appealcategory_category"
    df_m8 = spark.read.format("delta").load(m8_path)
    actual_m8_columns = df_m8.columns
    if set(expected_m8_columns) - set(actual_m8_columns):
        m8_test = False
        m8_diff = f"Missing columns in M8 output: {set(expected_m8_columns) - set(actual_m8_columns)}"

except Exception as e:
    print(f"Error during M1-M8 testing: {str(e)}")

# Define expected schemas
json_schema = StructType(
    [
        StructField("CaseNo", StringType(), True),
        StructField("HORef", StringType(), True),
        StructField("BailType", StringType(), True),
        StructField("Forenames", StringType(), True),
        StructField("Name", StringType(), True),
        StructField("BirthDate", StringType(), True),
        StructField("PortReference", StringType(), True),
        StructField("Detained", StringType(), True),
        StructField("FileLocationNote", StringType(), True),
        StructField("DecisionDate", StringType(), True),
        StructField("LegalHold", StringType(), True),
    ]
)

a360_schema = StructType(
    [
        StructField("operation", StringType(), True),
        StructField("relation_id", IntegerType(), True),
        StructField(
            "record_metadata",
            StructType(
                [
                    StructField("publisher", StringType(), True),
                    StructField("record_class", StringType(), True),
                    StructField("region", StringType(), True),
                    StructField("recordDate", StringType(), True),
                    StructField("event_date", StringType(), True),
                    StructField("client_identifier", StringType(), True),
                    StructField("bf_001", StringType(), True),  # HO Reference
                    StructField("bf_002", StringType(), True),  # Forenames
                    StructField("bf_003", StringType(), True),  # Surname
                    StructField("bf_004", StringType(), True),  # Date of Birth
                    StructField("bf_005", StringType(), True),  # Port Reference
                    StructField("bf_006", StringType(), True),  # Post code
                ]
            ),
            True,
        ),
        StructField(
            "file_metadata",
            StructType(
                [
                    StructField("publisher", StringType(), True),
                    StructField("dz_file_name", StringType(), True),
                    StructField("file_tag", StringType(), True),
                ]
            ),
            True,
        ),
    ]
)

# Read the output files
df_json = spark.read.json(f"{gold_mnt}/JSON/bails_*.json", schema=json_schema)
df_a360 = spark.read.json(f"{gold_mnt}/A360/bails_*.a360", schema=a360_schema)
df_html = spark.read.text(f"{gold_mnt}/HTML/bails_*.html")

# Initialise test result variables
json_record_count_test = True
a360_record_count_test = True
html_record_count_test = True
json_schema_test = True
a360_schema_test = True
json_record_count = 0
a360_record_count = 0
html_record_count = 0
json_schema_diff = ""
a360_schema_diff = ""

# Test output presence
try:
    json_record_count = spark.read.format("binaryFile").load(f"{gold_mnt}/JSON").count()
    assert json_record_count > 0, "No records found in JSON files"
except AssertionError as e:
    json_record_count_test = False

try:
    a360_record_count = spark.read.format("binaryFile").load(f"{gold_mnt}/A360").count()
    assert a360_record_count > 0, "No records found in A360 files"
except AssertionError as e:
    a360_record_count_test = False

try:
    html_record_count = spark.read.format("binaryFile").load(f"{gold_mnt}/HTML").count()
    assert html_record_count > 0, "No records found in HTML files"
except AssertionError as e:
    html_record_count_test = False

# Test schemas
expected_json_schema = json_schema
actual_json_schema = df_json.schema
json_schema_test = False

try:
    assert actual_json_schema == expected_json_schema, "JSON schema mismatch"
    json_schema_test = True
except AssertionError as e:
    json_schema_test = False
    json_schema_diff = f"Expected JSON Schema: {expected_json_schema}\nActual JSON Schema: {actual_json_schema}"

expected_a360_schema = a360_schema
actual_a360_schema = df_a360.schema
try:
    assert actual_a360_schema == expected_a360_schema, "A360 schema mismatch"
except AssertionError as e:
    a360_schema_test = False
    a360_schema_diff = f"Expected A360 Schema: {expected_a360_schema}\nActual A360 Schema: {actual_a360_schema}"

# Initialise transformation verification variables
segmentation_test = True
segmentation_diff = ""
metadata_test = True
metadata_diff = ""

# Test segmentation rules
expected_normal_conditions = [
    (col("BailType") == "2"),
    (~col("FileLocationNote").like("%destroyed%")),
    (~col("FileLocationNote").like("%detroyed%")),
    (~col("FileLocationNote").like("%destoyed%")),
    (~col("FileLocationNote").like("%distroyed%")),
    (to_date(col("DecisionDate"), "yyyy-MM-dd") > date_add(current_date(), -730))
]

normal_failures = df_json.filter(~reduce(lambda x, y: x & y, expected_normal_conditions)).count()

expected_legal_conditions = [(col("BailType") == "2"), (col("LegalHold") == "true")]

legal_failures = df_json.filter(~reduce(lambda x, y: x & y, expected_legal_conditions)).count()

if normal_failures > 0 or legal_failures > 0:
    segmentation_test = False
    segmentation_diff = f"Bails violations: {normal_failures} and {legal_failures}"

# Test metadata compliance
expected_metadata = {
    "publisher": "ARIA",
    "region": "GBR",
    "record_class": ["ARIAB", "ARIASB"],
}

metadata_violations = df_a360.filter(
    ~(col("record_metadata.publisher") == "ARIA")
    | ~(col("record_metadata.region") == "GBR")
    | ~col("record_metadata.record_class").isin(["ARIAB", "ARIASB"])
).count()

if metadata_violations > 0:
    metadata_test = False
    metadata_diff = f"Metadata violations found: {metadata_violations}"


def create_test_result_document():
    document = Document()

    # Add title
    document.add_heading("Bails Gold Output Testing Results", 0)

    # Introduction
    document.add_heading("Introduction", 1)
    document.add_paragraph("""This document presents the detailed test results for the gold output of the ARIA Bails data migration pipeline. 
    The tests cover bronze layer transformations (M1-M8) and gold output validation.""")

    # Bronze Layer Transformation Tests
    document.add_heading("Bronze Layer Transformation Tests", 1)

    document.add_heading("M1 transformation test", 2)
    document.add_paragraph(f"Expected columns in {m1_path} tested.")
    if m1_test == True:
        document.add_paragraph("M1 transformation test passed - all expected columns present in outputs.")
    else:
        document.add_paragraph("M1 transformation test failed - some expected columns missing in outputs.")
        document.add_paragraph(m1_diff)

    document.add_heading("M2 transformation test", 2)
    document.add_paragraph(f"Expected columns in {m2_path} tested.")
    if m2_test == True:
        document.add_paragraph("M2 transformation test passed - all expected columns present in outputs.")
    else:
        document.add_paragraph("M2 transformation test failed - some expected columns missing in outputs.")
        document.add_paragraph(m2_diff)

    document.add_heading("M3 transformation test", 2)
    document.add_paragraph(f"Expected columns in {m3_path} tested.")
    if m3_test == True:
        document.add_paragraph("M3 transformation test passed - all expected columns present in outputs.")
    else:
        document.add_paragraph("M3 transformation test failed - some expected columns missing in outputs.")
        document.add_paragraph(m3_diff)

    document.add_heading("M4 transformation test", 2)
    document.add_paragraph(f"Expected columns in {m4_path} tested.")
    if m4_test == True:
        document.add_paragraph("M4 transformation test passed - all expected columns present in outputs.")
    else:
        document.add_paragraph("M4 transformation test failed - some expected columns missing in outputs.")
        document.add_paragraph(m4_diff)

    document.add_heading("M5 transformation test", 2)
    document.add_paragraph(f"Expected columns in {m5_path} tested.")
    if m5_test == True:
        document.add_paragraph("M5 transformation test passed - all expected columns present in outputs.")
    else:
        document.add_paragraph("M5 transformation test failed - some expected columns missing in outputs.")
        document.add_paragraph(m5_diff)

    document.add_heading("M6 transformation test", 2)
    document.add_paragraph(f"Expected columns in {m6_path} tested.")
    if m6_test == True:
        document.add_paragraph("M6 transformation test passed - all expected columns present in outputs.")
    else:
        document.add_paragraph("M6 transformation test failed - some expected columns missing in outputs.")
        document.add_paragraph(m6_diff)

    document.add_heading("M7 transformation test", 2)
    document.add_paragraph(f"Expected columns in {m7_path} tested.")
    if m7_test == True:
        document.add_paragraph("M7 transformation test passed - all expected columns present in outputs.")
    else:
        document.add_paragraph("M7 transformation test failed - some expected columns missing in outputs.")
        document.add_paragraph(m7_diff)

    document.add_heading("M8 transformation test", 2)
    document.add_paragraph(f"Expected columns in {m8_path} tested.")
    if m8_test == True:
        document.add_paragraph("M8 transformation test passed - all expected columns present in outputs.")
    else:
        document.add_paragraph("M8 transformation test failed - some expected columns missing in outputs.")
        document.add_paragraph(m8_diff)

    # Gold Output Tests
    document.add_heading("Gold Output Tests", 1)

    # Data Presence and Format
    document.add_heading("Output Presence", 2)
    document.add_paragraph(f"JSON file count: {json_record_count}")
    document.add_paragraph(f"A360 file count: {a360_record_count}")
    document.add_paragraph(f"HTML file count: {html_record_count}")

    # Schema Validation
    document.add_heading("Schema Validation", 2)
    if json_schema_test:
        document.add_paragraph("JSON schema validation passed - all expected columns present in outputs.")
    else:
        document.add_paragraph("JSON schema validation failed - some expected columns missing in outputs.")
        document.add_paragraph(json_schema_diff)

    if a360_schema_test:
        document.add_paragraph("A360 schema validation passed - all expected columns present in outputs.")
    else:
        document.add_paragraph("A360 schema validation failed - some expected columns missing in outputs.")
        document.add_paragraph(a360_schema_diff)

    # Segmentation Rules
    document.add_heading("Segmentation Rules Validation", 2)
    if segmentation_test:
        document.add_paragraph("Segmentation rules validation passed.")
    else:
        document.add_paragraph("Segmentation rules validation failed.")
        document.add_paragraph(segmentation_diff)

    # Metadata Compliance
    document.add_heading("Metadata Compliance", 2)
    if metadata_test:
        document.add_paragraph("Metadata compliance validation passed.")
    else:
        document.add_paragraph("Metadata compliance validation failed.")
        document.add_paragraph(metadata_diff)

    # Data Quality Metrics
    document.add_heading("Data Quality Metrics", 1)

    # Conclusion
    document.add_heading("Conclusion", 1)
    all_tests_passed = all(
        [
            m1_test,
            m2_test,
            m3_test,
            m4_test,
            m5_test,
            m6_test,
            m7_test,
            m8_test,
            json_schema_test,
            a360_schema_test,
            segmentation_test,
            metadata_test,
        ]
    )

    if all_tests_passed:
        document.add_paragraph("""All validation tests have passed. Both bronze layer transformations 
        and gold output meet the expected criteria.""")
    else:
        document.add_paragraph("""Some validation tests have failed. Please review the detailed results 
        above for specific issues that need to be addressed.""")

    # Save the document
    document.save("bails_gold_output_test_results.docx")


# Execute tests and generate report
create_test_result_document()

# Print summary to console
print("Test Execution Complete")
print("Bronze Layer Tests:")
for i in range(1, 9):
    print(f"M{i} Test: {'Passed' if locals()[f'm{i}_test'] else 'Failed'}")
print("\nGold Output Tests:")
print(f"JSON Schema Test: {'Passed' if json_schema_test else 'Failed'}")
print(f"A360 Schema Test: {'Passed' if a360_schema_test else 'Failed'}")
print(f"Segmentation Test: {'Passed' if segmentation_test else 'Failed'}")
print(f"Metadata Test: {'Passed' if metadata_test else 'Failed'}")
