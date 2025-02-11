# %pip install azure-storage-blob pandas faker python-docx

from pyspark.sql.functions import col, count, when, sum, isnan, isnull
from pyspark.sql.types import (
    StructType,
    StructField,
    StringType,
    IntegerType,
    DateType,
    BooleanType,
    FloatType,
    TimestampType,
)

from docx import Document
from datetime import datetime

# Setting variables for use in subsequent cells
raw_mnt = "/mnt/raw/ARIADM/ARM/JOH"
landing_mnt = "/mnt/landing/test/"
bronze_mnt = "/mnt/bronze/ARIADM/ARM/JOH"
silver_mnt = "/mnt/silver/ARIADM/ARM/JOH"
gold_mnt = "/mnt/gold/ARIADM/ARM/JOH"

# Mapping abbreviations to human-readable text
abbreviation_map = {
    'adjudicator_et_hc_dnur': 'Adjudicator Employment Term, Hearing Centre, Do Not Use Reason',
    'johistory_users': 'Judicial Officer History, Users',
    'othercentre_hearingcentre': 'Other Centre Hearing Centre',
    'adjudicator_role': 'Adjudicator Role',
    'AdjudicatorId': 'Adjudicator ID',
    'SourceFileName': 'Source File Name',
    'JudicialStatus': 'Judicial Status',
    'missing_columns': 'Missing Columns',
    'data_type_mismatch_count': 'Data Type Mismatch Count',
    'exists': 'Exists',
    'count': 'Count',
    'null_count': 'Null Count'
}

def perform_adjudicator_schema_checks():
    """
    Performs key columns and schema consistency checks, designed for the Adjudicator DataFrame.

    Args:
        df (DataFrame): The DataFrame to perform checks on.
        table_name (str): The name of the table being checked.

    Returns:
        dict: A dictionary containing the results of the schema checks. The dictionary's keys are the table name and the value is a list of missing columns or a count of data type results.
    """
    schema_results = {}
    df = spark.read.format("delta").load(f"{bronze_mnt}/bronze_adjudicator_et_hc_dnur")
    table_name = "bronze_adjudicator_et_hc_dnur"

    # Check if all required columns are present
    required_columns = [
    'AdjudicatorId',
    'Surname',
    'Forenames',
    'Title',
    'DateOfBirth',
    'CorrespondenceAddress',
    'ContactTelephone',
    'ContactDetails',
    'AvailableAtShortNotice',
    'DesignatedCentre',
    'EmploymentTerm',
    'FullTime',
    'IdentityNumber',
    'DateOfRetirement',
    'ContractEndDate',
    'ContractRenewalDate',
    'DoNotUse',
    'DoNotUseReason',
    'JudicialStatus',
    'Address1',
    'Address2',
    'Address3',
    'Address4',
    'Address5',
    'Postcode',
    'Telephone',
    'Mobile',
    'Email',
    'BusinessAddress1',
    'BusinessAddress2',
    'BusinessAddress3',
    'BusinessAddress4',
    'BusinessAddress5',
    'BusinessPostcode',
    'BusinessTelephone',
    'BusinessFax',
    'BusinessEmail',
    'JudicialInstructions',
    'JudicialInstructionsDate',
    'Notes'
    ]
    missing_columns = [col for col in required_columns if col not in df.columns]
    schema_results[f"{table_name}_table_missing_columns"] = missing_columns

    # Check for data type consistency
    adjudicator_schema = StructType(
        [
            StructField("AdjudicatorId", IntegerType(), nullable=False),
            StructField("Surname", StringType(), nullable=False),
            StructField("Forenames", StringType(), nullable=False),
            StructField("Title", StringType(), nullable=False),
            StructField("DateOfBirth", DateType(), nullable=True),
            StructField("CorrespondenceAddress", IntegerType(), nullable=True),
            StructField("ContactTelephone", StringType(), nullable=True),
            StructField("ContactDetails", StringType(), nullable=True),
            StructField("AvailableAtShortNotice", BooleanType(), nullable=True),
            StructField("DesignatedCentre", StringType(), nullable=True),
            StructField("EmploymentTerm", StringType(), nullable=True),
            StructField("FullTime", IntegerType(), nullable=True),
            StructField("IdentityNumber", StringType(), nullable=True),
            StructField("DateOfRetirement", DateType(), nullable=True),
            StructField("ContractEndDate", DateType(), nullable=True),
            StructField("ContractRenewalDate", DateType(), nullable=True),
            StructField("DoNotUse", BooleanType(), nullable=True),
            StructField("DoNotUseReason", StringType(), nullable=True),
            StructField("JudicialStatus", IntegerType(), nullable=True),
            StructField("Address1", StringType(), nullable=True),
            StructField("Address2", StringType(), nullable=True),
            StructField("Address3", StringType(), nullable=True),
            StructField("Address4", StringType(), nullable=True),
            StructField("Address5", StringType(), nullable=True),
            StructField("Postcode", StringType(), nullable=True),
            StructField("Telephone", StringType(), nullable=True),
            StructField("Mobile", StringType(), nullable=True),
            StructField("Email", StringType(), nullable=True),
            StructField("BusinessAddress1", StringType(), nullable=True),
            StructField("BusinessAddress2", StringType(), nullable=True),
            StructField("BusinessAddress3", StringType(), nullable=True),
            StructField("BusinessAddress4", StringType(), nullable=True),
            StructField("BusinessAddress5", StringType(), nullable=True),
            StructField("BusinessPostcode", StringType(), nullable=True),
            StructField("BusinessTelephone", StringType(), nullable=True),
            StructField("BusinessFax", StringType(), nullable=True),
            StructField("BusinessEmail", StringType(), nullable=True),
            StructField("JudicialInstructions", StringType(), nullable=True),
            StructField("JudicialInstructionsDate", DateType(), nullable=True),
            StructField("Notes", StringType(), nullable=True),
            StructField("AdtclmnFirstCreatedDatetime", TimestampType(), nullable=True),
            StructField("AdtclmnModifiedDatetime", TimestampType(), nullable=True),
            StructField("SourceFileName", StringType(), nullable=True),
            StructField("InsertedByProcessName", StringType(), nullable=True)
        ]
    )

    data_type_mismatch_count = (
    df.select(
        [
        count(when(col(field.name).cast(field.dataType).isNull() & col(field.name).isNotNull(), 1)).alias(field.name)
        for field in adjudicator_schema.fields
        ]
        ).agg(*(sum(col(c)).alias(c) for c in df.columns))
        .collect()[0]
        )
    mismatch_count = 0
    for column in data_type_mismatch_count:
        mismatch_count += data_type_mismatch_count[column]
    schema_results[f"{table_name}_table_data_type_mismatch_count"] = mismatch_count

    return schema_results


def perform_johistory_users_schema_checks():
    """
    Performs key columns and schema consistency checks, designed for the JOHistory Users DataFrame.

    Args:
        df (DataFrame): The DataFrame to perform checks on.
        table_name (str): The name of the table being checked.

    Returns:
        dict: A dictionary containing the results of the schema checks. The dictionary's keys are the table name and the value is a list of missing columns or a count of data type results.
    """
    schema_results = {}
    df = spark.read.format("delta").load(f"{bronze_mnt}/bronze_johistory_users")
    table_name = "bronze_johistory_users"

    required_columns = [
    "AdjudicatorId",
    "HistDate",
    "HistType",
    "UserName",
    "Comment",
    "AdtclmnFirstCreatedDatetime",
    "AdtclmnModifiedDatetime",
    "SourceFileName",
    "InsertedByProcessName"
    ]
    missing_columns = [col for col in required_columns if col not in df.columns]
    schema_results[f"{table_name}_table_missing_columns"] = missing_columns

    # Check for data type consistency
    johistory_schema = StructType(
    [
        StructField("AdjudicatorId", IntegerType(), nullable=False),
        StructField("HistDate", DateType(), nullable=True),
        StructField("HistType", IntegerType(), nullable=True),
        StructField("UserName", StringType(), nullable=True),
        StructField("Comment", StringType(), nullable=True),
        StructField("AdtclmnFirstCreatedDatetime", TimestampType(), nullable=True),
        StructField("AdtclmnModifiedDatetime", TimestampType(), nullable=True),
        StructField("SourceFileName", StringType(), nullable=True),
        StructField("InsertedByProcessName", StringType(), nullable=True)
    ]
)

    data_type_mismatch_count = (
    df.select(
        [
        count(when(col(field.name).cast(field.dataType).isNull() & col(field.name).isNotNull(), 1)).alias(field.name)
        for field in johistory_schema.fields
        ]
        ).agg(*(sum(col(c)).alias(c) for c in df.columns))
        .collect()[0]
        )
    mismatch_count = 0
    for column in data_type_mismatch_count:
        mismatch_count += data_type_mismatch_count[column]
    schema_results[f"{table_name}_table_data_type_mismatch_count"] = mismatch_count

    return schema_results


def perform_othercentre_hearingcentre_schema_checks():
    """
    Performs key columns and schema consistency checks, designed for the Other Centre/Hearing Centre DataFrame.

    Args:
        df (DataFrame): The DataFrame to perform checks on.
        table_name (str): The name of the table being checked.

    Returns:
        dict: A dictionary containing the results of the schema checks. The dictionary's keys are the table name and the value is a list of missing columns or a count of data type results.
    """
    schema_results = {}
    df = spark.read.format("delta").load(f"{bronze_mnt}/bronze_othercentre_hearingcentre")
    table_name = "bronze_othercentre_hearingcentre"

    required_columns = [
    "AdjudicatorId",
    "OtherCentres",
    "AdtclmnFirstCreatedDatetime",
    "AdtclmnModifiedDatetime",
    "SourceFileName",
    "InsertedByProcessName"
    ]   
    missing_columns = [col for col in required_columns if col not in df.columns]
    schema_results[f"{table_name}_table_missing_columns"] = missing_columns

    # Check for data type consistency
    othercentre_hearingcentre_schema = StructType(
    [
        StructField("AdjudicatorId", IntegerType(), nullable=False),
        StructField("OtherCentres", StringType(), nullable=True),
        StructField("AdtclmnFirstCreatedDatetime", TimestampType(), nullable=True),
        StructField("AdtclmnModifiedDatetime", TimestampType(), nullable=True),
        StructField("SourceFileName", StringType(), nullable=True),
        StructField("InsertedByProcessName", StringType(), nullable=True)
    ]
)

    data_type_mismatch_count = (
    df.select(
        [
        count(when(col(field.name).cast(field.dataType).isNull() & col(field.name).isNotNull(), 1)).alias(field.name)
        for field in othercentre_hearingcentre_schema.fields
        ]
        ).agg(*(sum(col(c)).alias(c) for c in df.columns))
        .collect()[0]
        )
    mismatch_count = 0
    for column in data_type_mismatch_count:
        mismatch_count += data_type_mismatch_count[column]
    schema_results[f"{table_name}_table_data_type_mismatch_count"] = mismatch_count

    return schema_results


def perform_adjudicator_role_schema_checks():
    """
    Performs key columns and schema consistency checks, designed for the JOHistory Users DataFrame.

    Args:
        df (DataFrame): The DataFrame to perform checks on.
        table_name (str): The name of the table being checked.

    Returns:
        dict: A dictionary containing the results of the schema checks. The dictionary's keys are the table name and the value is a list of missing columns or a count of data type results.
    """
    schema_results = {}
    df = spark.read.format("delta").load(f"{bronze_mnt}/bronze_adjudicator_role")
    table_name =  "bronze_adjudicator_role"

    required_columns = [
    "AdjudicatorId",
    "Role",
    "DateOfAppointment",
    "EndDateOfAppointment",
    "AdtclmnFirstCreatedDatetime",
    "AdtclmnModifiedDatetime",
    "SourceFileName",
    "InsertedByProcessName"
]
    missing_columns = [col for col in required_columns if col not in df.columns]
    schema_results[f"{table_name}_table_missing_columns"] = missing_columns

    # Check for data type consistency
    adjudicator_role_schema = StructType(
    [
        StructField("AdjudicatorId", IntegerType(), nullable=False),
        StructField("Role", IntegerType(), nullable=True),
        StructField("DateOfAppointment", DateType(), nullable=True),
        StructField("EndDateOfAppointment", DateType(), nullable=True),
        StructField("AdtclmnFirstCreatedDatetime", TimestampType(), nullable=True),
        StructField("AdtclmnModifiedDatetime", TimestampType(), nullable=True),
        StructField("SourceFileName", StringType(), nullable=True),
        StructField("InsertedByProcessName", StringType(), nullable=True)
    ]
)

    data_type_mismatch_count = (
    df.select(
        [
        count(when(col(field.name).cast(field.dataType).isNull() & col(field.name).isNotNull(), 1)).alias(field.name)
        for field in adjudicator_role_schema.fields
        ]
        ).agg(*(sum(col(c)).alias(c) for c in df.columns))
        .collect()[0]
        )
    mismatch_count = 0
    for column in data_type_mismatch_count:
        mismatch_count += data_type_mismatch_count[column]
    schema_results[f"{table_name}_table_data_type_mismatch_count"] = mismatch_count

    return schema_results

def perform_data_quality_checks(df, table_name, key_column):
    """
    Performs data quality checks on the given DataFrame.

    Args:
        df (DataFrame): The DataFrame to perform checks on.
        table_name (str): The name of the table being checked.

    Returns:
        dict: A dictionary containing the validation results.
    """
    validation_results = {}

    # Check if the table exists
    validation_results[f"{table_name}_table_exists"] = df is not None

    # Check table count
    validation_results[f"{table_name}_row_count"] = df.count()

    # Check for null values in key columns
    key_columns = key_column
    for column in key_columns:
        if column in df.columns:
            null_count = df.filter(col(column).isNull()).count()
            validation_results[f"{table_name}_{column}_table_null_count"] = null_count

    return validation_results

def rename_dict_keys(input_dict):
    """
    Replaces known abbreviations with their full form readable format. 

    Args:
        input_dict (dict): The input dictionary with keys to be transformed.

    Returns:
        dict: A new dictionary with the transformed keys, preserving the original values.
    """
    renamed_dict = {}
    for key, value in input_dict.items():
        readable_key = key
        for abbr, readable in abbreviation_map.items():
            readable_key = readable_key.replace(abbr, readable)
        readable_key = readable_key.replace('_', ' ').capitalize()
        if readable_key.lower().startswith("bronze "):  # Remove "Bronze" if it exists
            readable_key = readable_key[7:]  # Skip the first 7 characters ("Bronze ")
        renamed_dict[readable_key] = value
    return renamed_dict


# Arranging the Bronze tables for data quality checks
# bronze_tables = [
#     ("bronze_adjudicator_et_hc_dnur", f"{bronze_mnt}/bronze_adjudicator_et_hc_dnur"),
#     ("bronze_johistory_users", f"{bronze_mnt}/bronze_johistory_users"),
#     ("bronze_othercentre_hearingcentre", f"{bronze_mnt}/bronze_othercentre_hearingcentre"),
#     ("bronze_adjudicator_role", f"{bronze_mnt}/bronze_adjudicator_role"),
# ]

# Optional - sanity check printout of Bronze tables
# df = spark.read.format("delta").load(f"{bronze_mnt}/bronze_adjudicator_et_hc_dnur")
# print('Adjudicator head:', df.head(), '\n')
# print('*******************************************')

# df = spark.read.format("delta").load(f"{bronze_mnt}/bronze_johistory_users")
# print('JOHistory head:', df.head(), '\n')
# print('*******************************************')

# df = spark.read.format("delta").load(f"{bronze_mnt}/bronze_othercentre_hearingcentre")
# print('Hearing centre head:', df.head(), '\n')
# print('*******************************************')

# df = spark.read.format("delta").load(f"{bronze_mnt}/bronze_adjudicator_role")
# print('Adjudicator role head:', df.head(), '\n')
print('<><><><><><><><><><><><><><><><><><><><><><>')

# Performing columns, schema and quality checks on data
validation_results = {}
validation_results.update(perform_data_quality_checks(spark.read.format("delta").load(f"{bronze_mnt}/bronze_adjudicator_et_hc_dnur"), "bronze_adjudicator_et_hc_dnur", ["AdjudicatorId", "JudicialStatus"]))
validation_results.update(perform_data_quality_checks(spark.read.format("delta").load(f"{bronze_mnt}/bronze_johistory_users"), "bronze_johistory_users", ["AdjudicatorId", "SourceFileName"]))
validation_results.update(perform_data_quality_checks(spark.read.format("delta").load(f"{bronze_mnt}/bronze_othercentre_hearingcentre"), "bronze_othercentre_hearingcentre", ["AdjudicatorId", "SourceFileName"]))
validation_results.update(perform_data_quality_checks(spark.read.format("delta").load(f"{bronze_mnt}/bronze_adjudicator_role"), "bronze_adjudicator_role", ["AdjudicatorId", "SourceFileName"]))

# print('Validation checks results: \n')
# for key,value in validation_results.items():
#     print(f'{key}: {value} \n')
# print('<><><><><><><><><><><><><><><><><><><><><><>')

schema_results = {}
schema_results.update(perform_adjudicator_schema_checks())
schema_results.update(perform_johistory_users_schema_checks())
schema_results.update(perform_othercentre_hearingcentre_schema_checks())
schema_results.update(perform_adjudicator_role_schema_checks())

# print('Schema checks results: \n')
# for key,value in schema_results.items():
#     print(f'{key}: {value} \n')
# print('<><><><><><><><><><><><><><><><><><><><><><>')

# Finalise results dictionary for output
overall_results = schema_results | validation_results
overall_results = rename_dict_keys(overall_results)

# Print the updated dictionary
print('Overall results: \n')
for key, value in overall_results.items():
    print(f'{key}: {value} \n')
print('<><><><><><><><><><><><><><><><><><><><><><> \n')

# Group metrics by section
grouped_results = {
    "Adjudicator Employment Term, Hearing Centre, Do Not Use Reason": [],
    "Judicial Officer History": [],
    "Other Centre Hearing Centre": [],
    "Adjudicator Role": []
}

# Categorize metrics into their respective groups
for metric, value in overall_results.items():
    if metric.lower().startswith("adjudicator employment term, hearing centre, do not use reason"):
        grouped_results["Adjudicator Employment Term, Hearing Centre, Do Not Use Reason"].append((metric, value))
    elif metric.lower().startswith("judicial officer history"):
        grouped_results["Judicial Officer History"].append((metric, value))
    elif metric.lower().startswith("other centre hearing centre"):
        grouped_results["Other Centre Hearing Centre"].append((metric, value))
    elif metric.lower().startswith("adjudicator role"):
        grouped_results["Adjudicator Role"].append((metric, value))

for key, value in grouped_results.items():
    print(f'{key}: {value} \n') 

# Create a new Word document
document = Document()

# Add a title to the document
document.add_heading('Bronze Data Quality Validation Report', level=1)

# Write grouped results into sections
for section, results in grouped_results.items():
    document.add_heading(f'{section} Test Results:', level=2)
    
    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Value'
    
    for metric, value in results:
        row_cells = table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = str(value)

# Add datetime stamp
document.add_paragraph('')
document.add_paragraph(f'Generated on {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')

# Save the document
document.save("bronze_data_quality_validation_report.docx")



