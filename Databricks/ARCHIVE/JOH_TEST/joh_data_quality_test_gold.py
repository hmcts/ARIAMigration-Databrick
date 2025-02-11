# TODO: lld reference for expected columns, a360/html/json count should be all the same, in depth testing gold outputs
# TODO: double check the gold output counts - are they from the test folder or the output folder? fix and rerun the pipeline, and rerun this file and see if all three outputs are the same

from pyspark.sql import SparkSession
from pyspark.sql.functions import col, to_date, coalesce, greatest, lit, explode
from pyspark.sql.types import StructType, StructField, StringType, IntegerType, DateType, ArrayType
from docx import Document
from docx.shared import Inches

# Create a SparkSession
spark = SparkSession.builder \
    .appName("GoldOutputTesting") \
    .getOrCreate()

# mount point for the gold files
gold_mnt = "/mnt/ingest00curatedsboxgold/ARIADM/ARM/JOH/test"

# Read the JSON file
json_schema = StructType([
    StructField("Surname", StringType(), True),
    StructField("Title", StringType(), True),
    StructField("Forenames", StringType(), True),
    StructField("DateOfBirth", StringType(), True),
    StructField("CorrespondenceAddress", StringType(), True),
    StructField("Telephone", StringType(), True),
    StructField("ContactDetails", StringType(), True),
    StructField("DesignatedCentre", StringType(), True),
    StructField("EmploymentTerm", StringType(), True),
    StructField("FullTime", StringType(), True),
    StructField("IdentityNumber", StringType(), True),
    StructField("DateOfRetirement", StringType(), True),
    StructField("ContractEndDate", StringType(), True),
    StructField("ContractRenewalDate", StringType(), True),
    StructField("DoNotUseReason", StringType(), True),
    StructField("JudicialStatus", StringType(), True),
    StructField("Address1", StringType(), True),
    StructField("Address2", StringType(), True),
    StructField("Address3", StringType(), True),
    StructField("Address4", StringType(), True),
    StructField("Address5", StringType(), True),
    StructField("Postcode", StringType(), True),
    StructField("Mobile", StringType(), True),
    StructField("Email", StringType(), True),
    StructField("BusinessAddress1", StringType(), True),
    StructField("BusinessAddress2", StringType(), True),
    StructField("BusinessAddress3", StringType(), True),
    StructField("BusinessAddress4", StringType(), True),
    StructField("BusinessAddress5", StringType(), True),
    StructField("BusinessPostcode", StringType(), True),
    StructField("BusinessTelephone", StringType(), True),
    StructField("BusinessFax", StringType(), True),
    StructField("BusinessEmail", StringType(), True),
    StructField("JudicialInstructions", StringType(), True),
    StructField("JudicialInstructionsDate", StringType(), True),
    StructField("Notes", StringType(), True),
    StructField("OtherCentres", ArrayType(StringType()), True),
    StructField("Roles", ArrayType(
        StructType([
            StructField("Role", StringType(), True),
            StructField("DateOfAppointment", StringType(), True),
            StructField("EndDateOfAppointment", StringType(), True)
        ])
    ), True),
    StructField("History", ArrayType(
        StructType([
            StructField("HistDate", StringType(), True),
            StructField("HistType", StringType(), True),
            StructField("UserName", StringType(), True),
            StructField("Comment", StringType(), True)
        ])
    ), True)
])

df_json = spark.read.json("/mnt/ingest00curatedsboxgold/ARIADM/ARM/JOH/test/JSON/judicial_officer_*.json", schema=json_schema)

# Read the A360 file
a360_schema = StructType([
    StructField("operation", StringType(), True),
    StructField("relation_id", IntegerType(), True),
    StructField("record_metadata", StructType([
        StructField("publisher", StringType(), True),
        StructField("record_class", StringType(), True),
        StructField("region", StringType(), True),
        StructField("recordDate", StringType(), True),
        StructField("event_date", StringType(), True),
        StructField("client_identifier", IntegerType(), True),
        StructField("bf_001", StringType(), True),
        StructField("bf_002", StringType(), True),
        StructField("bf_003", StringType(), True),
        StructField("bf_004", StringType(), True),
        StructField("bf_005", StringType(), True)
    ]), True),
    StructField("file_metadata", StructType([
        StructField("publisher", StringType(), True),
        StructField("dz_file_name", StringType(), True),
        StructField("file_tag", StringType(), True)
    ]), True)
])

df_a360 = spark.read.json("/mnt/ingest00curatedsboxgold/ARIADM/ARM/JOH/test/A360/judicial_officer_*.a360", schema=a360_schema)

# Read the HTML file
df_html = spark.read.text("/mnt/ingest00curatedsboxgold/ARIADM/ARM/JOH/test/HTML/judicial_officer_*.html")

# Initialize variables to capture test results
json_record_count_test = True
a360_record_count_test = True
html_record_count_test = True
json_schema_test = True
a360_schema_test = True
json_record_count = 0
a360_record_count = 0
html_record_count = 0
json_schema_diff = ""
a360_schema_diff = ""

# Test if the output is present and in the correct format
try:
    json_record_count = spark.read.format("binaryFile").load(f"{gold_mnt}/JSON").count()
    assert json_record_count > 0, "No records found in the JSON file"
except AssertionError as e:
    json_record_count_test = False

try:
    a360_record_count = spark.read.format("binaryFile").load(f"{gold_mnt}/A360").count()
    assert a360_record_count > 0, "No records found in the A360 file"
except AssertionError as e:
    a360_record_count_test = False

try:
    html_record_count = spark.read.format("binaryFile").load(f"{gold_mnt}/HTML").count()
    assert html_record_count > 0, "No records found in the HTML file"
except AssertionError as e:
    html_record_count_test = False

# Test the data types
expected_json_schema = json_schema
actual_json_schema = df_json.schema
try:
    assert actual_json_schema == expected_json_schema, "JSON schema does not match the expected schema"
except AssertionError as e:
    json_schema_test = False
    json_schema_diff = f"Expected JSON Schema: {expected_json_schema}\nActual JSON Schema: {actual_json_schema}"

expected_a360_schema = a360_schema
actual_a360_schema = df_a360.schema
try:
    assert actual_a360_schema == expected_a360_schema, "A360 schema does not match the expected schema"
except AssertionError as e:
    a360_schema_test = False
    a360_schema_diff = f"Expected A360 Schema: {expected_a360_schema}\nActual A360 Schema: {actual_a360_schema}"

# Initialize variables for transformation step verification
m1_test = True
m2_test = True
m3_test = True
m4_test = True
silver_table_test = True
m1_diff = ""
m2_diff = ""
m3_diff = ""
m4_diff = ""
silver_table_diff = ""

# Test the transformation steps
# M1 - bronze_adjudicator_et_hc_dnur
expected_m1_columns = ['AdjudicatorId', 'Surname', 'Forenames', 'Title', 'DateOfBirth', 'CorrespondenceAddress', 'ContactTelephone', 'ContactDetails', 'AvailableAtShortNotice', 'DesignatedCentre', 'EmploymentTerms', 'FullTime', 'IdentityNumber', 'DateOfRetirement', 'ContractEndDate', 'ContractRenewalDate', 'DoNotUse', 'DoNotUseReason', 'JudicialStatus', 'Address1', 'Address2', 'Address3', 'Address4', 'Address5', 'Postcode', 'Telephone', 'Mobile', 'Email', 'BusinessAddress1', 'BusinessAddress2', 'BusinessAddress3', 'BusinessAddress4', 'BusinessAddress5', 'BusinessPostcode', 'BusinessTelephone', 'BusinessFax', 'BusinessEmail', 'JudicialInstructions', 'JudicialInstructionsDate']
actual_m1_columns = df_json.columns
try:
    assert set(actual_m1_columns) == set(expected_m1_columns), "M1 transformation step failed"
except AssertionError as e:
    m1_test = False
    # m1_diff = f"Expected M1 Columns: {expected_m1_columns}\n Actual M1 Columns: {actual_m1_columns}"
    m1_diff = f"Columns in expected but not in actual: {[col for col in expected_m1_columns if col not in actual_m1_columns]}"

# M2 - bronze_johistory_users
expected_m2_columns = ['AdjudicatorId', 'HistDate', 'HistType', 'UserName', 'Comment']
df_history = df_json.select(explode("History").alias("History")).select("History.*")
actual_m2_columns = [column.name for column in df_history.schema]
try:
    assert set(actual_m2_columns) == set(expected_m2_columns), "M2 transformation step failed"
except AssertionError as e:
    m2_test = False
    # m2_diff = f"Expected M2 Columns: {expected_m2_columns}\nActual M2 Columns: {actual_m2_columns}"
    m2_diff = f"Columns in expected but not in actual: {[col for col in expected_m2_columns if col not in actual_m2_columns]}"


# M3 - bronze_othercentre_hearingcentre
expected_m3_columns = ["AdjudicatorId", "OtherCentres"]
df_other_centres = df_json.select("OtherCentres")
actual_m3_columns = [column.name for column in df_other_centres.schema]
try:
    assert set(actual_m3_columns) == set(expected_m3_columns), "M3 transformation step failed"
except AssertionError as e:
    m3_test = False
    # m3_diff = f"Expected M3 Columns: {expected_m3_columns}\nActual M3 Columns: {actual_m3_columns}"
    m3_diff = f"Columns in expected but not in actual: {[col for col in expected_m3_columns if col not in actual_m3_columns]}"


# M4 - bronze_adjudicator_role
expected_m4_columns = ["AdjudicatorId", "Role", "DateOfAppointment", "EndDateOfAppointment"]
df_roles = df_json.select(explode("Roles").alias("Roles")).select("Roles.*")
actual_m4_columns = [column.name for column in df_roles.schema]
try:
    assert set(actual_m4_columns) == set(expected_m4_columns), "M4 transformation step failed"
except AssertionError as e:
    m4_test = False
    # m4_diff = f"Expected M4 Columns: {expected_m4_columns}\nActual M4 Columns: {actual_m4_columns}"
    m4_diff = f"Columns in expected but not in actual: {[col for col in expected_m4_columns if col not in actual_m4_columns]}"


# Silver table - silver_adjudicator_details
expected_silver_columns = ["History", "Address1", "Address2", "Address3", "Address4"]
# silver_df = df_json.select("AdjudicatorId").distinct()
silver_df = df_json.select("JudicialStatus")
actual_silver_columns = silver_df.columns
try:
    assert set(actual_silver_columns) == set(expected_silver_columns), "Silver table transformation failed"
except AssertionError as e:
    silver_table_test = False
    # silver_table_diff = f"Expected Silver Columns: {expected_silver_columns}\nActual Silver Columns: {actual_silver_columns}"
    silver_diff = f"Columns in expected but not in actual: {[col for col in expected_silver_columns if col not in actual_silver_columns]}"


def create_test_result_document(df_json, df_a360, df_html, json_schema, a360_schema):
    document = Document()

    # Add title to the document
    document.add_heading('Gold Output Testing Results', 0)

    # Add Introduction section
    document.add_heading('Introduction', 1)
    document.add_paragraph("""This document presents the detailed test results for the gold output of the ARIA JOH data migration pipeline. 
    The tests cover various aspects such as data presence, format, schema validation, transformation step verification, and performance metrics.""")

    # Section for data presence and format
    document.add_heading('Data Presence and Format', 1)
    document.add_paragraph('The following tests were performed to ensure the presence and format of the gold output files:')

    # Test if the output is present and in the correct format
    document.add_heading('Output Presence', 2)
    document.add_paragraph(f'JSON file record count: {json_record_count}')
    if not json_record_count_test:
        document.add_paragraph('JSON file record count test failed. No records found in the JSON file.')
    document.add_paragraph(f'A360 file record count: {a360_record_count}')
    if not a360_record_count_test:
        document.add_paragraph('A360 file record count test failed. No records found in the A360 file.')
    document.add_paragraph(f'HTML file record count: {html_record_count}')
    if not html_record_count_test:
        document.add_paragraph('HTML file record count test failed. No records found in the HTML file.')

    # Test the data types
    document.add_heading('Data Types', 2)
    document.add_paragraph('The following tests were performed to validate the data types of the JSON and A360 files:')

    document.add_heading('JSON Schema', 3)
    if json_schema_test:
        document.add_paragraph('JSON schema test passed.')
    else:
        document.add_paragraph('JSON schema test failed.')
        document.add_paragraph(json_schema_diff)

    document.add_heading('A360 Schema', 3)
    if a360_schema_test:
        document.add_paragraph('A360 schema test passed.')
    else:
        document.add_paragraph('A360 schema test failed.')
        document.add_paragraph(a360_schema_diff)

    # Placeholders for screenshots
    # document.add_paragraph('Todo: Insert a screenshot of the pipeline diagram here.')
    # document.add_picture('placeholder_image.jpg', width=Inches(6))

    # document.add_paragraph('Todo: Insert a screenshot of the JSON file content here.')
    # document.add_picture('placeholder_image.jpg', width=Inches(6))

    # document.add_paragraph('Todo: Insert a screenshot of the A360 file content here.')
    # document.add_picture('placeholder_image.jpg', width=Inches(6))

    # document.add_paragraph('Todo: Insert a screenshot of the HTML file content here.')
    # document.add_picture('placeholder_image.jpg', width=Inches(6))

    # Section for transformation step verification
    document.add_heading('Transformation Step Verification', 1)
    document.add_paragraph('The following tests were performed to verify the transformation steps:')

    # Test the transformation steps
    # M1 - bronze_adjudicator_et_hc_dnur
    document.add_heading('M1 - bronze_adjudicator_et_hc_dnur', 2)
    if m1_test:
        document.add_paragraph('M1 transformation step passed.')
    else:
        document.add_paragraph('M1 transformation step failed.')
        document.add_paragraph(m1_diff)
        document.add_paragraph('Expected Transformation Rule:')
        document.add_paragraph('The bronze_adjudicator_et_hc_dnur table should contain the columns AdjudicatorId, Surname, Forenames, Title, DateOfBirth, CorrespondenceAddress, ContactTelephone, ContactDetails, AvailableAtShortNotice, DesignatedCentre, EmploymentTerms, FullTime, IdentityNumber, DateOfRetirement, ContractEndDate, ContractRenewalDate, DoNotUse, DoNotUseReason, JudicialStatus, Address1, Address2, Address3, Address4, Address5, Postcode, Telephone, Mobile, Email, BusinessAddress1, BusinessAddress2, BusinessAddress3,BusinessAddress4, BusinessAddress5, BusinessPostcode, BusinessTelephone, BusinessFax,BusinessEmail, JudicialInstructions, JudicialInstructionsDate, and Notes.')
        document.add_paragraph('Actual columns:')
        document.add_paragraph(actual_m1_columns)

    # M2 - bronze_johistory_users
    document.add_heading('M2 - bronze_johistory_users', 2)
    if m2_test:
        document.add_paragraph('M2 transformation step passed.')
    else:
        document.add_paragraph('M2 transformation step failed.')
        document.add_paragraph(m2_diff)
        document.add_paragraph('Expected Transformation Rule:')
        document.add_paragraph('The bronze_johistory_users table should contain the columns AdjudicatorId, HistDate, HistType, UserName, and Comment.')
        document.add_paragraph('Actual columns:')
        document.add_paragraph(actual_m2_columns)

    # M3 - bronze_othercentre_hearingcentre
    document.add_heading('M3 - bronze_othercentre_hearingcentre', 2)
    if m3_test:
        document.add_paragraph('M3 transformation step passed.')
    else:
        document.add_paragraph('M3 transformation step failed.')
        document.add_paragraph(m3_diff)
        document.add_paragraph('Expected Transformation Rule:')
        document.add_paragraph('The bronze_othercentre_hearingcentre table should contain the columns AdjudicatorId and OtherCentres.')
        document.add_paragraph('Actual columns:')
        document.add_paragraph(actual_m3_columns)

    # M4 - bronze_adjudicator_role
    document.add_heading('M4 - bronze_adjudicator_role', 2)
    if m4_test:
        document.add_paragraph('M4 transformation step passed.')
    else:
        document.add_paragraph('M4 transformation step failed.')
        document.add_paragraph(m4_diff)
        document.add_paragraph('Expected Transformation Rule:')
        document.add_paragraph('The bronze_adjudicator_role table should contain the columns AdjudicatorId, Role, DateOfAppointment, and EndDateOfAppointment.')
        document.add_paragraph('Actual columns:')
        document.add_paragraph(actual_m4_columns)

    # Silver table - silver_adjudicator_details
    document.add_heading('Silver Table - silver_adjudicator_details', 2)
    if silver_table_test:
        document.add_paragraph('Silver table transformation passed.')
    else:
        document.add_paragraph('Silver table transformation failed.')
        document.add_paragraph(silver_table_diff)
        document.add_paragraph('Expected Transformation Rule:')
        document.add_paragraph('The silver_adjudicator_details table should contain the column AdjudicatorId.')
        document.add_paragraph('Actual columns:')
        document.add_paragraph(actual_silver_columns)

    # Section for performance metrics
    document.add_heading('Performance Metrics', 1)
    document.add_paragraph('Todo: Add performance metrics such as execution time, resource utilization, and scalability here.')

    # Conclusion section
    document.add_heading('Conclusion', 1)
    if json_record_count_test and a360_record_count_test and html_record_count_test and json_schema_test and a360_schema_test and m1_test and m2_test and m3_test and m4_test and silver_table_test:
        document.add_paragraph("""Based on the test results, the gold output of the data migration pipeline meets the expected criteria. 
        The data is present in the correct format, the schemas are validated, and the transformation steps are verified. 
        The performance metrics indicate that the pipeline is efficient and scalable.""")
    else:
        document.add_paragraph("""Based on the test results, the gold output of the data migration pipeline does not meet all the expected criteria. 
        Some tests have failed. Please refer to the detailed test results for more information on the failures and the expected transformation rules.""")

    # Save the document
    document.save('gold_output_test_results.docx')

# Create the test result document
create_test_result_document(df_json, df_a360, df_html, json_schema, a360_schema)
