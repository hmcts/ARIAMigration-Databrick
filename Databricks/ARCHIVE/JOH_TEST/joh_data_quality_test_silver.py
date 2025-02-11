# %pip install azure-storage-blob pandas faker python-docx

from pyspark.sql.functions import col, count, when, sum
from pyspark.sql.types import (
    StructType, 
    StructField, 
    StringType, 
    IntegerType, 
    BooleanType, 
    DateType, 
    TimestampType,
    LongType
)

from docx import Document
from datetime import datetime

# Setting variables for use in subsequent cells
raw_mnt = "/mnt/raw/ARIADM/ARM/JOH"
landing_mnt = "/mnt/landing/test/"
bronze_mnt = "/mnt/bronze/ARIADM/ARM/JOH"
silver_mnt = "/mnt/silver/ARIADM/ARM/JOH"
gold_mnt = "/mnt/gold/ARIADM/ARM/JOH"

# Mapping abbreviations to human-readable text
abbreviation_map = {
    'adjudicator_details': 'Adjudicator Details',
    'adjudicator': 'Adjudicator',
    'role': 'Role',
    'AdjudicatorId': 'Adjudicator ID',
    'SourceFileName': 'Source File Name',
    'JudicialStatus': 'Judicial Status',
    'missing_columns': 'Missing Columns',
    'data_type_mismatch_count': 'Data Type Mismatch Count',
    'exists': 'Exists',
    'null_count': 'Null Count',
    'duplicate_count': 'Duplicate Count',
    'invalid_role_count': 'Invalid Role Count',
    'table': 'Table'
}

def perform_adjudicator_schema_checks():
    """
    Performs key columns and schema consistency checks, designed for the Silver Adjudicator DataFrame.

    Args:
        df (DataFrame): The DataFrame to perform checks on.
        table_name (str): The name of the table being checked.

    Returns:
        dict: A dictionary containing the results of the schema checks. The dictionary's keys are the table name and the value is a list of missing columns or a count of data type results.

    """
    schema_results = {}
    df = spark.read.format("delta").load(f"{silver_mnt}/silver_adjudicator_detail")
    table_name = "silver_adjudicator_detail"

    # Check if all required columns are present
    required_columns = [
    "AdjudicatorId", 
    "Surname", 
    "Forenames", 
    "Title", 
    "DateOfBirth", 
    "CorrespondenceAddress", 
    "ContactDetails", 
    "ContactTelephone", 
    "AvailableAtShortNotice", 
    "DesignatedCentre", 
    "EmploymentTerm", 
    "FullTime", 
    "IdentityNumber", 
    "DateOfRetirement", 
    "ContractEndDate", 
    "ContractRenewalDate", 
    "DoNotUse", 
    "DoNotUseReason", 
    "JudicialStatus", 
    "Address1", 
    "Address2", 
    "Address3", 
    "Address4", 
    "Address5", 
    "Postcode", 
    "Telephone", 
    "Mobile", 
    "Email", 
    "BusinessAddress1", 
    "BusinessAddress2", 
    "BusinessAddress3", 
    "BusinessAddress4", 
    "BusinessAddress5", 
    "BusinessPostcode", 
    "BusinessTelephone", 
    "BusinessFax", 
    "BusinessEmail", 
    "JudicialInstructions", 
    "JudicialInstructionsDate", 
    "Notes", 
    "AdtclmnFirstCreatedDatetime", 
    "AdtclmnModifiedDatetime", 
    "SourceFileName", 
    "InsertedByProcessName"
    ]
    missing_columns = [col for col in required_columns if col not in df.columns]
    schema_results[f"{table_name}_table_missing_columns"] = missing_columns

    # Check for data type consistency
    schema = StructType(
        [
            StructField("AdjudicatorId", IntegerType(), nullable=False),
            StructField("Surname", StringType(), nullable=False),
            StructField("Forenames", StringType(), nullable=False),
            StructField("Title", StringType(), nullable=False),
            StructField("DateOfBirth", DateType(), nullable=True),
            StructField("CorrespondenceAddress", StringType(), nullable=True),
            StructField("ContactDetails", StringType(), nullable=True),
            StructField("ContactTelephone", StringType(), nullable=True),
            StructField("AvailableAtShortNotice", BooleanType(), nullable=True),
            StructField("DesignatedCentre", StringType(), nullable=True),
            StructField("EmploymentTerm", StringType(), nullable=True),
            StructField("FullTime", StringType(), nullable=True),
            StructField("IdentityNumber", IntegerType(), nullable=True),
            StructField("DateOfRetirement", DateType(), nullable=True),
            StructField("ContractEndDate", DateType(), nullable=True),
            StructField("ContractRenewalDate", DateType(), nullable=True),
            StructField("DoNotUse", BooleanType(), nullable=True),
            StructField("DoNotUseReason", StringType(), nullable=True),
            StructField("JudicialStatus", StringType(), nullable=True),
            StructField("Address1", StringType(), nullable=True),
            StructField("Address2", StringType(), nullable=True),
            StructField("Address3", StringType(), nullable=True),
            StructField("Address4", StringType(), nullable=True),
            StructField("Address5", StringType(), nullable=True),
            StructField("Postcode", StringType(), nullable=True),
            StructField("Telephone", StringType(), nullable=True),
            StructField("Mobile", StringType(), nullable=True),
            StructField("Email", StringType(), nullable=True),
            StructField("BusinessAddress1", StringType(), nullable=True),
            StructField("BusinessAddress2", StringType(), nullable=True),
            StructField("BusinessAddress3", StringType(), nullable=True),
            StructField("BusinessAddress4", StringType(), nullable=True),
            StructField("BusinessAddress5", StringType(), nullable=True),
            StructField("BusinessPostcode", StringType(), nullable=True),
            StructField("BusinessTelephone", StringType(), nullable=True),
            StructField("BusinessFax", StringType(), nullable=True),
            StructField("BusinessEmail", StringType(), nullable=True),
            StructField("JudicialInstructions", StringType(), nullable=True),
            StructField("JudicialInstructionsDate", DateType(), nullable=True),
            StructField("Notes", StringType(), nullable=True),
            StructField("AdtclmnFirstCreatedDatetime", TimestampType(), nullable=True),
            StructField("AdtclmnModifiedDatetime", TimestampType(), nullable=True),
            StructField("SourceFileName", StringType(), nullable=True),
            StructField("InsertedByProcessName", StringType(), nullable=True)
        ]
    )
    data_type_mismatch_count = (
        df.select(
        [
        count(when(col(field.name).cast(field.dataType).isNull() & col(field.name).isNotNull(), 1)).alias(field.name)
        for field in schema.fields
        ]
        ).agg(*(sum(col(c)).alias(c) for c in df.columns))
        .collect()[0]
        )
    mismatch_count = 0
    for column in data_type_mismatch_count:
        mismatch_count += data_type_mismatch_count[column]
    schema_results[f"{table_name}_table_data_type_mismatch_count"] = mismatch_count

    return schema_results

def perform_history_detail_schema_checks():
    """
    Performs key columns and schema consistency checks, designed for the Silver History Detail DataFrame.

    Args:
        df (DataFrame): The DataFrame to perform checks on.
        table_name (str): The name of the table being checked.

    Returns:
        dict: A dictionary containing the results of the schema checks. The dictionary's keys are the table name and the value is a list of missing columns or a count of data type results.

    """
    schema_results = {}
    df = spark.read.format("delta").load(f"{silver_mnt}/silver_history_detail")
    table_name = "silver_history_detail"

    # Check if all required columns are present
    required_columns = [
    "AdjudicatorId",
    "HistDate",
    "HistType",
    "UserName",
    "Comment",
    "AdtclmnFirstCreatedDatetime",
    "AdtclmnModifiedDatetime",
    "SourceFileName",
    "InsertedByProcessName"
    ]
    missing_columns = [col for col in required_columns if col not in df.columns]
    schema_results[f"{table_name}_table_missing_columns"] = missing_columns

    # Check for data type consistency
    schema = StructType([
        StructField("AdjudicatorId", LongType(), True),
        StructField("HistDate", StringType(), True),
        StructField("HistType", StringType(), True),
        StructField("UserName", StringType(), True),
        StructField("Comment", StringType(), True),
        StructField("AdtclmnFirstCreatedDatetime", TimestampType(), True),
        StructField("AdtclmnModifiedDatetime", TimestampType(), True),
        StructField("SourceFileName", StringType(), True),
        StructField("InsertedByProcessName", StringType(), True)
    ])
    data_type_mismatch_count = (
        df.select(
        [
        count(when(col(field.name).cast(field.dataType).isNull() & col(field.name).isNotNull(), 1)).alias(field.name)
        for field in schema.fields
        ]
        ).agg(*(sum(col(c)).alias(c) for c in df.columns))
        .collect()[0]
        )
    mismatch_count = 0
    for column in data_type_mismatch_count:
        mismatch_count += data_type_mismatch_count[column]
    schema_results[f"{table_name}_table_data_type_mismatch_count"] = mismatch_count

    return schema_results

def perform_appointment_detail_schema_checks():
    """
    Performs key columns and schema consistency checks, designed for the Silver Appointment Detail DataFrame.

    Args:
        df (DataFrame): The DataFrame to perform checks on.
        table_name (str): The name of the table being checked.

    Returns:
        dict: A dictionary containing the results of the schema checks. The dictionary's keys are the table name and the value is a list of missing columns or a count of data type results.

    """
    schema_results = {}
    df = spark.read.format("delta").load(f"{silver_mnt}/silver_appointment_detail")
    table_name = "silver_appointment_detail"

    # Check if all required columns are present
    required_columns = [
    "AdjudicatorId",
    "Role",
    "DateOfAppointment",
    "EndDateOfAppointment",
    "AdtclmnFirstCreatedDatetime",
    "AdtclmnModifiedDatetime",
    "SourceFileName",
    "InsertedByProcessName"
    ]
    missing_columns = [col for col in required_columns if col not in df.columns]
    schema_results[f"{table_name}_table_missing_columns"] = missing_columns

    # Check for data type consistency
    schema = StructType([
        StructField("AdjudicatorId", LongType(), True),
        StructField("Role", StringType(), True),
        StructField("DateOfAppointment", StringType(), True),
        StructField("EndDateOfAppointment", StringType(), True),
        StructField("AdtclmnFirstCreatedDatetime", TimestampType(), True),
        StructField("AdtclmnModifiedDatetime", TimestampType(), True),
        StructField("SourceFileName", StringType(), True),
        StructField("InsertedByProcessName", StringType(), True)
    ])
    data_type_mismatch_count = (
        df.select(
        [
        count(when(col(field.name).cast(field.dataType).isNull() & col(field.name).isNotNull(), 1)).alias(field.name)
        for field in schema.fields
        ]
        ).agg(*(sum(col(c)).alias(c) for c in df.columns))
        .collect()[0]
        )
    mismatch_count = 0
    for column in data_type_mismatch_count:
        mismatch_count += data_type_mismatch_count[column]
    schema_results[f"{table_name}_table_data_type_mismatch_count"] = mismatch_count

    return schema_results

def perform_othercentre_detail_schema_checks():
    """
    Performs key columns and schema consistency checks, designed for the Silver Other Centre Detail DataFrame.

    Args:
        df (DataFrame): The DataFrame to perform checks on.
        table_name (str): The name of the table being checked.

    Returns:
        dict: A dictionary containing the results of the schema checks. The dictionary's keys are the table name and the value is a list of missing columns or a count of data type results.

    """
    schema_results = {}
    df = spark.read.format("delta").load(f"{silver_mnt}/silver_othercentre_detail")
    table_name = "silver_othercentre_detail"

    # Check if all required columns are present
    required_columns = [
    "AdjudicatorId",
    "OtherCentres",
    "AdtclmnFirstCreatedDatetime",
    "AdtclmnModifiedDatetime",
    "SourceFileName",
    "InsertedByProcessName"
    ]
    missing_columns = [col for col in required_columns if col not in df.columns]
    schema_results[f"{table_name}_table_missing_columns"] = missing_columns

    # Check for data type consistency
    schema = StructType([
        StructField("AdjudicatorId", LongType(), True),
        StructField("OtherCentres", StringType(), True),
        StructField("AdtclmnFirstCreatedDatetime", TimestampType(), True),
        StructField("AdtclmnModifiedDatetime", TimestampType(), True),
        StructField("SourceFileName", StringType(), True),
        StructField("InsertedByProcessName", StringType(), True)
    ])
    data_type_mismatch_count = (
        df.select(
        [
        count(when(col(field.name).cast(field.dataType).isNull() & col(field.name).isNotNull(), 1)).alias(field.name)
        for field in schema.fields
        ]
        ).agg(*(sum(col(c)).alias(c) for c in df.columns))
        .collect()[0]
        )
    mismatch_count = 0
    for column in data_type_mismatch_count:
        mismatch_count += data_type_mismatch_count[column]
    schema_results[f"{table_name}_table_data_type_mismatch_count"] = mismatch_count

    return schema_results

def perform_data_quality_checks(df, table_name, key_column):
    """
    Performs data quality checks on the given DataFrame.

    Args:
        df (DataFrame): The DataFrame to perform checks on.
        table_name (str): The name of the table being checked.

    Returns:
        dict: A dictionary containing the validation results.
    """
    validation_results = {}

    # Check if the table exists
    validation_results[f"{table_name}_table_exists"] = df is not None

    # Check table count
    validation_results[f"{table_name}_row_count"] = df.count()

    # Check for null values in key columns
    key_columns = key_column
    for column in key_columns:
        if column in df.columns:
            null_count = df.filter(col(column).isNull()).count()
            validation_results[f"{table_name}_{column}_table_null_count"] = null_count

    return validation_results

def rename_dict_keys(input_dict):
    """
    Replaces known abbreviations with their full form readable format. 

    Args:
        input_dict (dict): The input dictionary with keys to be transformed.

    Returns:
        dict: A new dictionary with the transformed keys, preserving the original values.
    """
    renamed_dict = {}
    for key, value in input_dict.items():
        readable_key = key
        for abbr, readable in abbreviation_map.items():
            readable_key = readable_key.replace(abbr, readable)
        readable_key = readable_key.replace('_', ' ').capitalize()
        if readable_key.lower().startswith("silver "):  # Remove "Silver" if it exists
            readable_key = readable_key[7:]  # Skip the first 7 characters ("Silver ")
        renamed_dict[readable_key] = value
    return renamed_dict

# Executing checks
schema_results = {}
schema_results.update(perform_adjudicator_schema_checks())
schema_results.update(perform_history_detail_schema_checks())
schema_results.update(perform_appointment_detail_schema_checks())
schema_results.update(perform_othercentre_detail_schema_checks())

validation_results = {}
validation_results.update(perform_data_quality_checks(spark.read.format("delta").load(f"{silver_mnt}/silver_adjudicator_detail"), "silver_adjudicator_detail", ["AdjudicatorId", "JudicialStatus"]))
validation_results.update(perform_data_quality_checks(spark.read.format("delta").load(f"{silver_mnt}/silver_appointment_detail"), "silver_appointment_detail", ["AdjudicatorId", "SourceFileName"]))
validation_results.update(perform_data_quality_checks(spark.read.format("delta").load(f"{silver_mnt}/silver_history_detail"), "silver_history_detail", ["AdjudicatorId", "SourceFileName"]))
validation_results.update(perform_data_quality_checks(spark.read.format("delta").load(f"{silver_mnt}/silver_othercentre_detail"), "silver_othercentre_detail", ["AdjudicatorId", "SourceFileName"]))

# Gathering results
overall_results = schema_results | validation_results
overall_results = rename_dict_keys(overall_results)

# Print the updated dictionary
print('Overall results: \n')
for key, value in overall_results.items():
    print(f'{key}: {value} \n')
print('<><><><><><><><><><><><><><><><><><><><><><>')

# Group metrics by section
grouped_results = {
    "Adjudicator Detail": [],
    "History Detail": [],
    "Appointment Detail": [],
    "Other Centre Detail": []
}

# Categorize metrics into their respective groups
for metric, value in overall_results.items():
    if metric.lower().startswith("adjudicator detail"):
        grouped_results["Adjudicator Detail"].append((metric, value))
    elif metric.lower().startswith("history detail"):
        grouped_results["History Detail"].append((metric, value))
    elif metric.lower().startswith("appointment detail"):
        grouped_results["Appointment Detail"].append((metric, value))
    elif metric.lower().startswith("othercentre detail"):
        grouped_results["Other Centre Detail"].append((metric, value))

for key, value in grouped_results.items():
    print(f'{key}: {value} \n') 

# Create a new Word document
document = Document()

# Add a title to the document
document.add_heading("Silver Data Quality Validation Report", level=1)

# Write grouped results into sections
for section, results in grouped_results.items():
    document.add_heading(f'{section} Test Results:', level=2)
    
    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Value'
    
    for metric, value in results:
        row_cells = table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = str(value)

# Add datetime stamp
document.add_paragraph('')
document.add_paragraph(f'Generated on {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')

# Save the document
document.save("silver_data_quality_validation_report.docx")
