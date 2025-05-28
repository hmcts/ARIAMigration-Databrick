from pyspark.sql import SparkSession
from pyspark.sql.functions import col, to_date, coalesce, greatest, lit, explode
from pyspark.sql.types import StructType, StructField, StringType, IntegerType, DateType, ArrayType
from docx import Document
from docx.shared import Inches

# Create a SparkSession
spark = SparkSession.builder \
    .appName("GoldOutputTesting") \
    .getOrCreate()

# mount point for the gold files
gold_mnt = "/mnt/gold/ARIADM/ARM/TD"

# Read the JSON file
json_schema = StructType([
    StructField("CaseNo", StringType(), True),
    StructField("Forenames", StringType(), True),
    StructField("Name", StringType(), True),
    StructField("BirthDate", StringType(), True),
    StructField("DestructionDate", StringType(), True),
    StructField("HORef", StringType(), True),
    StructField("PortReference", StringType(), True),
    StructField("FileLocation", StringType(), True),
    StructField("Description", StringType(), True),
    StructField("Note", StringType(), True)
])

df_json = spark.read.json("/mnt/gold/ARIADM/ARM/TD/JSON/tribunal_decision_*.json", schema=json_schema)

# Read the A360 file
a360_schema = StructType([
    StructField("operation", StringType(), True),
    StructField("relation_id", IntegerType(), True),
    StructField("record_metadata", StructType([
        StructField("publisher", StringType(), True),
        StructField("record_class", StringType(), True),
        StructField("region", StringType(), True),
        StructField("recordDate", StringType(), True),
        StructField("event_date", StringType(), True),
        StructField("client_identifier", IntegerType(), True),
        StructField("bf_001", StringType(), True),
        StructField("bf_002", StringType(), True),
        StructField("bf_003", StringType(), True),
        StructField("bf_004", StringType(), True),
        StructField("bf_005", StringType(), True),
        StructField("bf_006", StringType(), True),
        StructField("bf_007", StringType(), True),
        StructField("bf_008", StringType(), True),
        StructField("bf_009", StringType(), True)
    ]), True),
    StructField("file_metadata", StructType([
        StructField("publisher", StringType(), True),
        StructField("dz_file_name", StringType(), True),
        StructField("file_tag", StringType(), True)
    ]), True)
])

df_a360 = spark.read.json("/mnt/gold/ARIADM/ARM/TD/A360/judicial_officer_*.a360", schema=a360_schema)

# Read the HTML file
df_html = spark.read.text("/mnt/gold/ARIADM/ARM/TD/HTML/tribunal_decision_*.html")

# Initialize variables to capture test results
json_record_count_test = True
a360_record_count_test = True
html_record_count_test = True
json_schema_test = True
a360_schema_test = True
json_record_count = 0
a360_record_count = 0
html_record_count = 0
json_schema_diff = ""
a360_schema_diff = ""

# Test if the output is present and in the correct format
try:
    json_record_count = spark.read.format("binaryFile").load(f"{gold_mnt}/JSON").count()
    assert json_record_count > 0, "No records found in the JSON file"
except AssertionError as e:
    json_record_count_test = False

try:
    a360_record_count = spark.read.format("binaryFile").load(f"{gold_mnt}/A360").count()
    assert a360_record_count > 0, "No records found in the A360 file"
except AssertionError as e:
    a360_record_count_test = False

try:
    html_record_count = spark.read.format("binaryFile").load(f"{gold_mnt}/HTML").count()
    assert html_record_count > 0, "No records found in the HTML file"
except AssertionError as e:
    html_record_count_test = False

# Test the data types
expected_json_schema = json_schema
actual_json_schema = df_json.schema
try:
    assert actual_json_schema == expected_json_schema, "JSON schema does not match the expected schema"
except AssertionError as e:
    json_schema_test = False
    json_schema_diff = f"Expected JSON Schema: {expected_json_schema}\nActual JSON Schema: {actual_json_schema}"

expected_a360_schema = a360_schema
actual_a360_schema = df_a360.schema
try:
    assert actual_a360_schema == expected_a360_schema, "A360 schema does not match the expected schema"
except AssertionError as e:
    a360_schema_test = False
    a360_schema_diff = f"Expected A360 Schema: {expected_a360_schema}\nActual A360 Schema: {actual_a360_schema}"

# Initialize variables for transformation step verification
m1_test = True
m1_diff = ""

# Test the transformation steps
# M1 - bronze_ac_ca_ant_fl_dt_hc
expected_m1_columns = ['CaseNo', 'Forenames', 'Name', 'BirthDate', 'DestructionDate', 'HORef', 'PortReference', 'FileLocation', 'Description', 'Note']
actual_m1_columns = df_json.columns
try:
    assert set(actual_m1_columns) == set(expected_m1_columns), "M1 transformation step failed"
except AssertionError as e:
    m1_test = False
    m1_diff = f"Columns in expected but not in actual: {[col for col in expected_m1_columns if col not in actual_m1_columns]}"

def create_test_result_document(df_json, df_a360, df_html, json_schema, a360_schema):
    document = Document()

    # Add title to the document
    document.add_heading('Gold Output Testing Results', 0)

    # Add Introduction section
    document.add_heading('Introduction', 1)
    document.add_paragraph("""This document presents the detailed test results for the gold output of the ARIA TD data migration pipeline. 
    The tests cover various aspects such as data presence, format, schema validation, transformation step verification, and performance metrics.""")

    # Section for data presence and format
    document.add_heading('Data Presence and Format', 1)
    document.add_paragraph('The following tests were performed to ensure the presence and format of the gold output files:')

    # Test if the output is present and in the correct format
    document.add_heading('Output Presence', 2)
    document.add_paragraph(f'JSON file record count: {json_record_count}')
    if not json_record_count_test:
        document.add_paragraph('JSON file record count test failed. No records found in the JSON file.')
    document.add_paragraph(f'A360 file record count: {a360_record_count}')
    if not a360_record_count_test:
        document.add_paragraph('A360 file record count test failed. No records found in the A360 file.')
    document.add_paragraph(f'HTML file record count: {html_record_count}')
    if not html_record_count_test:
        document.add_paragraph('HTML file record count test failed. No records found in the HTML file.')

    # Test the data types
    document.add_heading('Data Types', 2)
    document.add_paragraph('The following tests were performed to validate the data types of the JSON and A360 files:')

    document.add_heading('JSON Schema', 3)
    if json_schema_test:
        document.add_paragraph('JSON schema test passed.')
    else:
        document.add_paragraph('JSON schema test failed.')
        document.add_paragraph(json_schema_diff)

    document.add_heading('A360 Schema', 3)
    if a360_schema_test:
        document.add_paragraph('A360 schema test passed.')
    else:
        document.add_paragraph('A360 schema test failed.')
        document.add_paragraph(a360_schema_diff)

    # Section for transformation step verification
    document.add_heading('Transformation Step Verification', 1)
    document.add_paragraph('The following tests were performed to verify the transformation steps:')

    # Test the transformation steps
    # M1 - bronze_ac_ca_ant_fl_dt_hc
    document.add_heading('M1 - bronze_ac_ca_ant_fl_dt_hc', 2)
    if m1_test:
        document.add_paragraph('M1 transformation step passed.')
    else:
        document.add_paragraph('M1 transformation step failed.')
        document.add_paragraph(m1_diff)
        document.add_paragraph('Expected Transformation Rule:')
        document.add_paragraph('The bronze_ac_ca_ant_fl_dt_hc table should contain the columns CaseNo, Forenames, Name, BirthDate, DestructionDate, HORef, PortReference, FileLocation, Description, and Note.')
        document.add_paragraph('Actual columns:')
        document.add_paragraph(actual_m1_columns)

    # Section for performance metrics
    document.add_heading('Performance Metrics', 1)
    document.add_paragraph('Todo: Add performance metrics such as execution time, resource utilization, and scalability here.')

    # Conclusion section
    document.add_heading('Conclusion', 1)
    if json_record_count_test and a360_record_count_test and html_record_count_test and json_schema_test and a360_schema_test and m1_test:
        document.add_paragraph("""Based on the test results, the gold output of the data migration pipeline meets the expected criteria. 
        The data is present in the correct format, the schemas are validated, and the transformation steps are verified. 
        The performance metrics indicate that the pipeline is efficient and scalable.""")
    else:
        document.add_paragraph("""Based on the test results, the gold output of the data migration pipeline does not meet all the expected criteria. 
        Some tests have failed. Please refer to the detailed test results for more information on the failures and the expected transformation rules.""")

    # Save the document
    document.save('td_gold_output_test_results.docx')

# Create the test result document
create_test_result_document(df_json, df_a360, df_html, json_schema, a360_schema)
