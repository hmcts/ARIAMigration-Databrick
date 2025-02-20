import json
from pyspark.sql import SparkSession
from pyspark.sql.functions import col, to_date, coalesce, greatest, lit, explode
from pyspark.sql.types import (
    StructType,
    StructField,
    StringType,
    IntegerType,
    DateType,
    ArrayType,
)
from pyspark.sql.functions import (
    col,
    count,
    avg,
    when,
    sum,
    current_date,
    datediff,
    to_date,
)
from docx import Document
from docx.shared import Inches
from datetime import datetime

# Create a SparkSession
spark = SparkSession.builder.appName("AppealsGoldOutputTesting").getOrCreate()

# Mount points
gold_mnt = "/mnt/gold/ARIADM/ARM/Appeals"

# Define schemas based on LLD document
appeal_case_schema = StructType(
    [
        StructField("CaseNo", StringType(), True),
        StructField("CasePrefix", StringType(), True),
        StructField("CaseYear", IntegerType(), True),
        StructField("CaseType", StringType(), True),
        StructField("AppealTypeId", IntegerType(), True),
        StructField("DateLodged", DateType(), True),
        StructField("DateReceived", DateType(), True),
        StructField("PortId", IntegerType(), True),
        StructField("HORef", StringType(), True),
        StructField("DateServed", DateType(), True),
        StructField("Notes", StringType(), True),
        StructField("NationalityId", IntegerType(), True),
        StructField("Interpreter", StringType(), True),
        StructField("CountryId", IntegerType(), True),
        StructField("DateOfIssue", DateType(), True),
        StructField("FamilyCase", StringType(), True),
        StructField("OakingtonCase", StringType(), True),
        StructField("VisitVisaType", StringType(), True),
        StructField("HOInterpreter", StringType(), True),
        StructField("AdditionalGrounds", StringType(), True),
        StructField("AppealCategories", StringType(), True),
        StructField("DateApplicationLodged", DateType(), True),
        StructField("ThirdCountryId", IntegerType(), True),
        StructField("StatutoryClosureDate", DateType(), True),
        StructField("PubliclyFunded", StringType(), True),
        StructField("NonStandardSCPeriod", StringType(), True),
        StructField("CourtPreference", StringType(), True),
        StructField("ProvisionalDestructionDate", DateType(), True),
        StructField("DestructionDate", DateType(), True),
        StructField("FileInStatutoryClosure", StringType(), True),
        StructField("DateOfNextListedHearing", DateType(), True),
        StructField("DocumentsReceived", StringType(), True),
        StructField("OutOfTimeIssue", StringType(), True),
        StructField("ValidityIssues", StringType(), True),
        StructField("ReceivedFromRespondent", StringType(), True),
        StructField("DateAppealReceived", DateType(), True),
        StructField("RemovalDate", DateType(), True),
        StructField("CaseOutcomeId", IntegerType(), True),
        StructField("AppealReceivedBy", StringType(), True),
        StructField("InCamera", StringType(), True),
        StructField("SecureCourtRequired", StringType(), True),
        StructField("DateOfApplicationDecision", DateType(), True),
        StructField("UserId", IntegerType(), True),
        StructField("SubmissionURN", StringType(), True),
        StructField("DateReinstated", DateType(), True),
        StructField("DeportationDate", DateType(), True),
        StructField("HOANRef", StringType(), True),
        StructField("HumanRights", StringType(), True),
        StructField("TransferOutDate", DateType(), True),
        StructField("CertifiedDate", DateType(), True),
        StructField("CertifiedRecordedDate", DateType(), True),
        StructField("NoticeSentDate", DateType(), True),
        StructField("AddressRecordedDate", DateType(), True),
        StructField("ReferredToJudgeDate", DateType(), True),
        StructField("CentreId", IntegerType(), True),
        StructField("FeeSatisfactionId", IntegerType(), True),
    ]
)

# Read the data files
try:
    df_appeal_case = spark.read.parquet(f"{gold_mnt}/AppealCase/*.parquet")
    df_case_appellant = spark.read.parquet(f"{gold_mnt}/CaseAppellant/*.parquet")
    df_list = spark.read.parquet(f"{gold_mnt}/List/*.parquet")
    df_bf_diary = spark.read.parquet(f"{gold_mnt}/BFDiary/*.parquet")
    df_history = spark.read.parquet(f"{gold_mnt}/History/*.parquet")
    df_link = spark.read.parquet(f"{gold_mnt}/Link/*.parquet")
    df_status = spark.read.parquet(f"{gold_mnt}/Status/*.parquet")
    df_appeal_category = spark.read.parquet(f"{gold_mnt}/AppealCategory/*.parquet")
    df_transaction = spark.read.parquet(f"{gold_mnt}/Transaction/*.parquet")
    df_human_right = spark.read.parquet(f"{gold_mnt}/HumanRight/*.parquet")
    df_new_matter = spark.read.parquet(f"{gold_mnt}/NewMatter/*.parquet")
    df_documents = spark.read.parquet(f"{gold_mnt}/Documents/*.parquet")
    df_direction = spark.read.parquet(f"{gold_mnt}/Direction/*.parquet")
except Exception as e:
    print(f"Error reading parquet files: {str(e)}")

# Variables to capture test results
appeal_case_test = True
case_appellant_test = True
list_test = True
bf_diary_test = True
history_test = True
link_test = True
status_test = True
appeal_category_test = True
transaction_test = True
human_right_test = True
new_matter_test = True
documents_test = True
direction_test = True


"""
Tests based on LLD
Case segmentation rules
Payment calculation rules
Status flow rules
Required field validations
Date format validations
"""


def test_bronze_data_import():
    """Test Bronze layer data import for M1-M16"""
    try:
        results = {"status": True, "errors": [], "metrics": {}}

        # Test M1-M16 data presence
        for module in range(1, 17):
            table_count = spark.sql(f"SELECT COUNT(*) FROM bronze_m{module}").collect()[
                0
            ][0]
            results["metrics"][f"m{module}_record_count"] = table_count
            if table_count == 0:
                results["status"] = False
                results["errors"].append(f"No data found in M{module} bronze table")

        # Test individual bronze tables presence
        additional_tables = [
            "bronze_cost_award",
            "bronze_cost_order",
            "bronze_appeal_type_category",
            "bronze_review_specific_direction",
            "bronze_hearing_point_change_reason",
            "bronze_hearing_point_history",
        ]

        for table in additional_tables:
            table_count = spark.sql(f"SELECT COUNT(*) FROM {table}").collect()[0][0]
            results["metrics"][f"{table}_record_count"] = table_count
            if table_count == 0:
                results["status"] = False
                results["errors"].append(f"No data found in {table}")

        return results
    except Exception as e:
        return {
            "status": False,
            "errors": [f"Bronze data import test failed with error: {str(e)}"],
            "metrics": {},
        }


def test_silver_transformations():
    """Test Silver layer transformations"""
    try:
        results = {"status": True, "errors": [], "metrics": {}}

        # Test appeal segmentations
        segmentation_tests = {
            "ft_appeals": "DateOfApplicationDecision IS NOT NULL AND DATEDIFF(current_date(), DateOfApplicationDecision) <= 180",
            "ut_appeals": "DateOfApplicationDecision IS NOT NULL AND DATEDIFF(current_date(), DateOfApplicationDecision) <= 1825",
            "skeleton_cases": "CasePrefix IN ('IA','LD','LE','LH','LP','LR') AND HOANRef IS NOT NULL",
        }

        for seg_name, condition in segmentation_tests.items():
            count = spark.sql(
                f"SELECT COUNT(*) FROM silver_appeals WHERE {condition}"
            ).collect()[0][0]
            results["metrics"][f"{seg_name}_count"] = count
            if count == 0:
                results["status"] = False
                results["errors"].append(
                    f"No records found for segmentation {seg_name}"
                )

        # Test metadata generation
        metadata_fields = [
            "client_identifier",
            "event_date",
            "recordDate",
            "region",
            "publisher",
            "record_class",
            "entitlement_tag",
        ]

        for field in metadata_fields:
            null_count = spark.sql(
                f"SELECT COUNT(*) FROM silver_appeals WHERE {field} IS NULL"
            ).collect()[0][0]
            if null_count > 0:
                results["status"] = False
                results["errors"].append(
                    f"Found {null_count} records with null {field}"
                )

        # Test business metadata generation
        business_metadata = {
            "bf_001": "HORef",
            "bf_002": "Forenames",
            "bf_003": "Name",
            "bf_004": "BirthDate",
            "bf_005": "PortReference",
            "bf_006": "Postcode",
        }

        for bf_field, source_field in business_metadata.items():
            null_count = spark.sql(
                f"SELECT COUNT(*) FROM silver_appeals WHERE {bf_field} IS NULL"
            ).collect()[0][0]
            if null_count > 0:
                results["status"] = False
                results["errors"].append(
                    f"Found {null_count} records with null {bf_field}"
                )

        return results
    except Exception as e:
        return {
            "status": False,
            "errors": [f"Silver transformations test failed with error: {str(e)}"],
            "metrics": {},
        }


def test_gold_outputs():
    """Test Gold layer outputs"""
    try:
        results = {"status": True, "errors": [], "metrics": {}}

        # Test HTML files
        html_files = spark.read.format("binaryFile").load(f"{gold_mnt}/HTML")
        results["metrics"]["html_file_count"] = html_files.count()

        # Test JSON files
        json_files = spark.read.format("binaryFile").load(f"{gold_mnt}/JSON")
        results["metrics"]["json_file_count"] = json_files.count()

        # Test manifest files
        manifest_files = spark.read.format("binaryFile").load(f"{gold_mnt}/manifest")
        results["metrics"]["manifest_file_count"] = manifest_files.count()

        # Test batch size (250 records per batch)
        manifest_records = spark.read.json(f"{gold_mnt}/manifest/*.json")
        batch_sizes = manifest_records.groupBy("batch_id").count()
        invalid_batches = batch_sizes.filter("count > 250").count()
        if invalid_batches > 0:
            results["status"] = False
            results["errors"].append(
                f"Found {invalid_batches} batches exceeding 250 records limit"
            )

        # Test file naming uniqueness
        duplicate_html = html_files.groupBy("path").count().filter("count > 1").count()
        duplicate_json = json_files.groupBy("path").count().filter("count > 1").count()
        duplicate_manifest = (
            manifest_files.groupBy("path").count().filter("count > 1").count()
        )

        if any([duplicate_html, duplicate_json, duplicate_manifest]):
            results["status"] = False
            results["errors"].append("Found duplicate file names")

        return results
    except Exception as e:
        return {
            "status": False,
            "errors": [f"Gold outputs test failed with error: {str(e)}"],
            "metrics": {},
        }


def test_file_formats():
    """Test output file formats"""
    try:
        results = {"status": True, "errors": [], "metrics": {}}

        # HTML Format Validation
        html_files = spark.read.format("binaryFile").load(f"{gold_mnt}/HTML/*.html")
        for row in html_files.collect():
            if not row.content.decode("utf-8").strip().startswith("<!DOCTYPE html>"):
                results["status"] = False
                results["errors"].append(f"Invalid HTML format in file: {row.path}")

        # JSON Format Validation
        json_files = spark.read.format("binaryFile").load(f"{gold_mnt}/JSON/*.json")
        for row in json_files.collect():
            try:
                json.loads(row.content.decode("utf-8"))
            except json.JSONDecodeError:
                results["status"] = False
                results["errors"].append(f"Invalid JSON format in file: {row.path}")

        # Manifest Format Validation
        manifest_files = spark.read.format("binaryFile").load(
            f"{gold_mnt}/manifest/*.json"
        )
        for row in manifest_files.collect():
            try:
                manifest = json.loads(row.content.decode("utf-8"))
                if not all(
                    key in manifest for key in ["batch_id", "records", "timestamp"]
                ):
                    results["status"] = False
                    results["errors"].append(
                        f"Missing required fields in manifest file: {row.path}"
                    )
            except json.JSONDecodeError:
                results["status"] = False
                results["errors"].append(f"Invalid manifest format in file: {row.path}")

        return results
    except Exception as e:
        return {
            "status": False,
            "errors": [f"File format test failed with error: {str(e)}"],
            "metrics": {},
        }


def test_batch_size():
    """Test manifest batch size compliance"""
    try:
        results = {"status": True, "errors": [], "metrics": {}}

        manifest_files = spark.read.json(f"{gold_mnt}/manifest/*.json")
        batch_sizes = manifest_files.groupBy("batch_id").agg(
            count("*").alias("record_count")
        )

        # Check for oversized batches
        oversized = batch_sizes.filter(col("record_count") > 250)
        if oversized.count() > 0:
            results["status"] = False
            results["errors"].append(
                f"Found {oversized.count()} batches exceeding 250 records"
            )

        # Check for undersized batches (except last batch)
        max_batch_id = batch_sizes.agg(max("batch_id")).collect()[0][0]
        undersized = batch_sizes.filter(
            (col("record_count") < 250) & (col("batch_id") != max_batch_id)
        )
        if undersized.count() > 0:
            results["status"] = False
            results["errors"].append(f"Found {undersized.count()} undersized batches")

        results["metrics"]["total_batches"] = batch_sizes.count()
        results["metrics"]["avg_batch_size"] = batch_sizes.agg(
            avg("record_count")
        ).collect()[0][0]

        return results
    except Exception as e:
        return {
            "status": False,
            "errors": [f"Batch size test failed with error: {str(e)}"],
            "metrics": {},
        }


def test_date_formats(df, date_columns):
    """Test if date columns are in correct format"""
    try:
        for date_col in date_columns:
            # Try converting to date format
            invalid_dates = df.filter(
                col(date_col).isNotNull() & to_date(col(date_col)).isNull()
            ).count()

            if invalid_dates > 0:
                return False
        return True
    except Exception:
        return False


def test_required_fields(df, required_fields):
    """Test if required fields are present and not null"""
    try:
        for field in required_fields:
            if field not in df.columns:
                return False

            null_count = df.filter(col(field).isNull()).count()
            if null_count > 0:
                return False
        return True
    except Exception:
        return False


def test_appeal_case(df_appeal_case):
    """Test Appeal Case data quality"""
    try:
        results = {"status": True, "errors": []}

        # Date columns validation
        date_columns = [
            "DateLodged",
            "DateReceived",
            "DateServed",
            "DateOfIssue",
            "DateApplicationLodged",
            "StatutoryClosureDate",
            "ProvisionalDestructionDate",
            "DestructionDate",
            "DateOfNextListedHearing",
            "DateAppealReceived",
            "RemovalDate",
            "DateOfApplicationDecision",
            "DateReinstated",
            "DeportationDate",
            "TransferOutDate",
            "CertifiedDate",
            "CertifiedRecordedDate",
            "NoticeSentDate",
            "AddressRecordedDate",
            "ReferredToJudgeDate",
        ]

        if not test_date_formats(df_appeal_case, date_columns):
            results["status"] = False
            results["errors"].append("Invalid date formats found")

        # Required fields validation
        required_fields = ["CaseNo", "CasePrefix", "CaseYear", "CaseType"]
        if not test_required_fields(df_appeal_case, required_fields):
            results["status"] = False
            results["errors"].append("Missing required fields")

        # Case number format validation
        invalid_case_numbers = df_appeal_case.filter(
            ~col("CaseNo").rlike("^[A-Z]{2}[0-9]{6}/[0-9]{4}$")
        ).count()
        if invalid_case_numbers > 0:
            results["status"] = False
            results["errors"].append(
                f"Invalid case number format found in {invalid_case_numbers} records"
            )

        return results
    except Exception as e:
        return {
            "status": False,
            "errors": [f"Appeal Case test failed with error: {str(e)}"],
        }


def test_case_appellant(df_case_appellant):
    """Test Case Appellant data quality"""
    try:
        results = {"status": True, "errors": []}

        required_fields = ["AppellantId", "CaseNo", "Relationship"]
        if not test_required_fields(df_case_appellant, required_fields):
            results["status"] = False
            results["errors"].append("Missing required fields")

        invalid_dob = df_case_appellant.filter(
            (col("BirthDate").isNotNull())
            & ((col("BirthDate") < "1900-01-01") | (col("BirthDate") > current_date()))
        ).count()

        if invalid_dob > 0:
            results["status"] = False
            results["errors"].append(
                f"Invalid date of birth found in {invalid_dob} records"
            )

        return results
    except Exception as e:
        return {
            "status": False,
            "errors": [f"Case Appellant test failed with error: {str(e)}"],
        }


def test_transactions(df_transaction):
    """Test Transaction data quality"""
    try:
        results = {"status": True, "errors": []}

        required_fields = ["TransactionId", "CaseNo", "Amount"]
        if not test_required_fields(df_transaction, required_fields):
            results["status"] = False
            results["errors"].append("Missing required fields")

        payment_summary = df_transaction.groupBy("CaseNo").agg(
            sum(when(col("TransactionTypeId") == 1, col("Amount")).otherwise(0)).alias(
                "FirstTierFee"
            ),
            sum(when(col("SumFeeAdjustment") == 1, col("Amount")).otherwise(0)).alias(
                "TotalFeeAdjustments"
            ),
            sum(when(col("SumTotalFee") == 1, col("Amount")).otherwise(0)).alias(
                "TotalFeeDue"
            ),
            sum(when(col("SumTotalPay") == 1, abs(col("Amount"))).otherwise(0)).alias(
                "TotalPaymentsReceived"
            ),
            sum(when(col("SumPayAdjustment") == 1, col("Amount")).otherwise(0)).alias(
                "TotalPaymentAdjustments"
            ),
            sum(when(col("SumBalance") == 1, col("Amount")).otherwise(0)).alias(
                "BalanceDue"
            ),
        )

        invalid_payments = payment_summary.filter(
            (col("TotalFeeDue") != (col("FirstTierFee") + col("TotalFeeAdjustments")))
            | (
                col("BalanceDue")
                != (
                    col("TotalFeeDue")
                    - col("TotalPaymentsReceived")
                    - col("TotalPaymentAdjustments")
                )
            )
        ).count()

        if invalid_payments > 0:
            results["status"] = False
            results["errors"].append(
                f"Invalid payment calculations found in {invalid_payments} records"
            )

        return results
    except Exception as e:
        return {
            "status": False,
            "errors": [f"Transaction test failed with error: {str(e)}"],
        }


def test_status(df_status):
    """Test Status data quality and flow"""
    try:
        results = {"status": True, "errors": []}

        # Required fields validation
        required_fields = ["StatusId", "CaseNo", "CaseStatus"]
        if not test_required_fields(df_status, required_fields):
            results["status"] = False
            results["errors"].append("Missing required fields in Status")

        # Status flow validation
        invalid_status_flow = df_status.filter(
            (
                col("CaseStatus").isin(
                    [40, 41, 42, 43, 44, 45, 53, 27, 28, 29, 34, 32, 33]
                )
            )
            & (col("Outcome") == 86)
            & (~col("PreviousStatus").isin([37, 38, 39, 17]))
        ).count()

        if invalid_status_flow > 0:
            results["status"] = False
            results["errors"].append(
                f"Invalid status flow found in {invalid_status_flow} records"
            )

        return results
    except Exception as e:
        return {"status": False, "errors": [f"Status test failed with error: {str(e)}"]}


def test_list(df_list):
    """Test List data quality"""
    try:
        results = {"status": True, "errors": []}

        # Required fields validation
        required_fields = ["StatusId", "CaseNo", "ListNumber", "HearingDuration"]
        if not test_required_fields(df_list, required_fields):
            results["status"] = False
            results["errors"].append("Missing required fields in List")

        # Hearing duration format validation
        invalid_duration = df_list.filter(
            (col("HearingDuration").isNotNull())
            & (~col("HearingDuration").rlike("^([0-9]{1,2}):([0-5][0-9])$"))
        ).count()

        if invalid_duration > 0:
            results["status"] = False
            results["errors"].append(
                f"Invalid hearing duration format in {invalid_duration} records"
            )

        return results
    except Exception as e:
        return {"status": False, "errors": [f"List test failed with error: {str(e)}"]}


def test_human_rights(df_appeal_case, df_human_right):
    """Test Human Rights data quality and consistency"""
    try:
        results = {"status": True, "errors": []}

        # Required fields validation
        required_fields = ["CaseNo", "HumanRightId"]
        if not test_required_fields(df_human_right, required_fields):
            results["status"] = False
            results["errors"].append("Missing required fields in Human Rights")

        # Flag consistency validation
        hr_inconsistency = (
            df_appeal_case.join(df_human_right, "CaseNo", "left")
            .filter((col("HumanRights") == "1") & (col("HumanRightId").isNull()))
            .count()
        )

        if hr_inconsistency > 0:
            results["status"] = False
            results["errors"].append(
                f"Human Rights flag inconsistency found in {hr_inconsistency} records"
            )

        return results
    except Exception as e:
        return {
            "status": False,
            "errors": [f"Human Rights test failed with error: {str(e)}"],
        }


def test_documents(df_documents):
    """Test Documents data quality"""
    try:
        results = {"status": True, "errors": []}

        # Required fields validation
        required_fields = ["CaseNo", "ReceivedDocumentId", "DateReceived"]
        if not test_required_fields(df_documents, required_fields):
            results["status"] = False
            results["errors"].append("Missing required fields in Documents")

        # Document dates validation
        invalid_doc_dates = df_documents.filter(
            (col("DateReceived").isNotNull())
            & (col("DateRequired").isNotNull())
            & (col("DateReceived") < col("DateRequired"))
        ).count()

        if invalid_doc_dates > 0:
            results["status"] = False
            results["errors"].append(
                f"Invalid document dates found in {invalid_doc_dates} records"
            )

        return results
    except Exception as e:
        return {
            "status": False,
            "errors": [f"Documents test failed with error: {str(e)}"],
        }


def test_direction(df_direction):
    """Test Direction data quality"""
    try:
        results = {"status": True, "errors": []}

        # Required fields validation
        required_fields = ["CaseNo", "StatusId"]
        if not test_required_fields(df_direction, required_fields):
            results["status"] = False
            results["errors"].append("Missing required fields in Direction")

        # Direction dates validation
        invalid_direction_dates = df_direction.filter(
            (col("DateRequiredIND").isNotNull())
            & (col("DateReceivedIND").isNotNull())
            & (col("DateReceivedIND") < col("DateRequiredIND"))
        ).count()

        if invalid_direction_dates > 0:
            results["status"] = False
            results["errors"].append(
                f"Invalid direction dates found in {invalid_direction_dates} records"
            )

        return results
    except Exception as e:
        return {
            "status": False,
            "errors": [f"Direction test failed with error: {str(e)}"],
        }


def test_metadata_fields():
    """Test metadata fields format and values"""
    try:
        results = {"status": True, "errors": [], "metrics": {}}

        # Archive metadata validation
        archive_metadata_rules = {
            "client_identifier": lambda x: x.rlike("^[A-Z]{2}[0-9]{6}/[0-9]{4}$"),
            "event_date": lambda x: x.cast("timestamp").isNotNull(),
            "recordDate": lambda x: x.cast("timestamp").isNotNull(),
            "region": lambda x: x == "GBR",
            "publisher": lambda x: x == "ARIA",
            "record_class": lambda x: x.isin(["ARIAFTA", "ARIAUTA"]),
            "entitlement_tag": lambda x: x == "IA_Tribunal",
        }

        for field, validation_rule in archive_metadata_rules.items():
            invalid_count = df_appeal_case.filter(~validation_rule(col(field))).count()
            if invalid_count > 0:
                results["status"] = False
                results["errors"].append(
                    f"Found {invalid_count} records with invalid {field}"
                )

        # Business metadata validation
        business_metadata_rules = {
            "bf_001": lambda x: x.isNotNull(),  # HO Reference
            "bf_002": lambda x: x.isNotNull(),  # Forenames
            "bf_003": lambda x: x.isNotNull(),  # Surname
            "bf_004": lambda x: x.cast("date").isNotNull(),  # Date of Birth
            "bf_005": lambda x: x.isNotNull(),  # Port Reference
            "bf_006": lambda x: x.rlike(
                "^[A-Z]{1,2}[0-9][A-Z0-9]? ?[0-9][A-Z]{2}$"
            ),  # Postcode
        }

        for field, validation_rule in business_metadata_rules.items():
            invalid_count = df_appeal_case.filter(~validation_rule(col(field))).count()
            if invalid_count > 0:
                results["status"] = False
                results["errors"].append(
                    f"Found {invalid_count} records with invalid {field}"
                )

        return results
    except Exception as e:
        return {
            "status": False,
            "errors": [f"Metadata field test failed with error: {str(e)}"],
            "metrics": {},
        }


def test_segmentation(df_appeal_case):
    """Test case segmentation rules"""
    try:
        results = {"status": True, "errors": [], "metrics": {}}

        # FT Retained - ARM
        ft_retained_arm = df_appeal_case.filter(
            (col("CaseType") == "1")
            & (col("DateOfApplicationDecision").isNotNull())
            & (datediff(current_date(), col("DateOfApplicationDecision")) <= 730)
        ).count()

        results["metrics"]["ft_retained_arm"] = ft_retained_arm

        # UT Retained
        ut_retained = df_appeal_case.filter(
            (col("CaseType") == "1")
            & (col("DateOfApplicationDecision").isNotNull())
            & (datediff(current_date(), col("DateOfApplicationDecision")) <= 1825)
        ).count()

        results["metrics"]["ut_retained"] = ut_retained

        # Skeleton Cases
        skeleton_cases = df_appeal_case.filter(
            col("CasePrefix").isin(["IA", "LD", "LE", "LH", "LP", "LR"])
            & col("HOANRef").isNotNull()
        ).count()

        results["metrics"]["skeleton_cases"] = skeleton_cases

        # Validate segmentation rules
        if ft_retained_arm == 0:
            results["status"] = False
            results["errors"].append("No First Tier Retained ARM cases found")

        if ut_retained == 0:
            results["status"] = False
            results["errors"].append("No Upper Tribunal Retained cases found")

        return results
    except Exception as e:
        return {
            "status": False,
            "errors": [f"Segmentation test failed with error: {str(e)}"],
            "metrics": {},
        }


def test_business_rules(df_appeal_case, df_status):
    """Test business rules compliance"""
    try:
        results = {"status": True, "errors": []}

        # Rule 1: First Tier retention period (2 years)
        ft_retention_violation = df_appeal_case.filter(
            (col("record_class") == "ARIAFTA")
            & (col("DateOfApplicationDecision").isNotNull())
            & (datediff(current_date(), col("DateOfApplicationDecision")) > 730)
        ).count()

        if ft_retention_violation > 0:
            results["status"] = False
            results["errors"].append(
                f"First Tier retention period violated in {ft_retention_violation} cases"
            )

        # Rule 2: Upper Tribunal retention period (5 years)
        ut_retention_violation = df_appeal_case.filter(
            (col("record_class") == "ARIAUTA")
            & (col("DateOfApplicationDecision").isNotNull())
            & (datediff(current_date(), col("DateOfApplicationDecision")) > 1825)
        ).count()

        if ut_retention_violation > 0:
            results["status"] = False
            results["errors"].append(
                f"Upper Tribunal retention period violated in {ut_retention_violation} cases"
            )

        # Rule 3: Status Flow Validation
        invalid_status_flow = (
            df_status.join(
                df_status.alias("prev_status"),
                (df_status.CaseNo == col("prev_status.CaseNo"))
                & (df_status.StatusId > col("prev_status.StatusId")),
            )
            .filter(
                ~(
                    col("CaseStatus").isin(
                        [
                            "40",
                            "41",
                            "42",
                            "43",
                            "44",
                            "45",
                            "53",
                            "27",
                            "28",
                            "29",
                            "34",
                            "32",
                            "33",
                        ]
                    )
                    & col("prev_status.CaseStatus").isin(["37", "38", "39", "17"])
                )
            )
            .count()
        )

        if invalid_status_flow > 0:
            results["status"] = False
            results["errors"].append(
                f"Invalid status flow found in {invalid_status_flow} cases"
            )

        return results
    except Exception as e:
        return {
            "status": False,
            "errors": [f"Business rules test failed with error: {str(e)}"],
        }


def generate_test_report(test_results):
    """Generate comprehensive test report"""
    try:
        document = Document()

        # Title
        document.add_heading("Appeals Archive Gold Output - Test Results", 0)

        # Executive Summary
        document.add_heading("Executive Summary", 1)
        all_passed = all(result["status"] for result in test_results.values())

        if all_passed:
            document.add_paragraph("✓ All validation tests passed successfully")
        else:
            document.add_paragraph(
                "✗ Some validation tests failed - see detailed results below"
            )

        # Detailed Results
        document.add_heading("Detailed Test Results", 1)

        for test_name, result in test_results.items():
            document.add_heading(test_name.replace("_", " ").title(), 2)

            status = "✓ Passed" if result["status"] else "✗ Failed"
            document.add_paragraph(f"Status: {status}")

            if not result["status"]:
                document.add_paragraph("Issues found:")
                for error in result["errors"]:
                    document.add_paragraph(f"• {error}", style="List Bullet")

            # Add metrics if available
            if "metrics" in result:
                document.add_paragraph("Metrics:")
                for metric_name, value in result["metrics"].items():
                    document.add_paragraph(
                        f"• {metric_name}: {value}", style="List Bullet"
                    )

        # Save the report
        document.save("appeals_gold_output_test_results.docx")
        return True
    except Exception as e:
        print(f"Error generating test report: {str(e)}")
        return False


def main():
    """Main execution flow"""
    try:
        # Initialize test results dictionary
        test_results = {}

        # Execute all tests
        # Bronze layer tests
        test_results["bronze_import"] = test_bronze_data_import()
        # Silver layer tests
        test_results["silver_transformations"] = test_silver_transformations()
        # Gold layer tests
        test_results["gold_outputs"] = test_gold_outputs()

        # File format tests
        test_results["file_formats"] = test_file_formats()
        test_results["batch_size"] = test_batch_size()
        test_results["metadata_fields"] = test_metadata_fields()

        # other more detailed testing
        test_results["appeal_case"] = test_appeal_case(df_appeal_case)
        test_results["case_appellant"] = test_case_appellant(df_case_appellant)
        test_results["transactions"] = test_transactions(df_transaction)
        test_results["status"] = test_status(df_status)
        test_results["list"] = test_list(df_list)
        test_results["human_rights"] = test_human_rights(df_appeal_case, df_human_right)
        test_results["documents"] = test_documents(df_documents)
        test_results["direction"] = test_direction(df_direction)
        test_results["metadata"] = test_metadata_fields(df_appeal_case)
        test_results["segmentation"] = test_segmentation(df_appeal_case)
        test_results["business_rules"] = test_business_rules(df_appeal_case, df_status)

        # Generate test report
        report_generated = generate_test_report(test_results)

        if report_generated:
            print("Test execution completed. Report generated successfully.")
        else:
            print("Test execution completed but report generation failed.")

        # Return overall test status
        return all(result["status"] for result in test_results.values())

    except Exception as e:
        print(f"Error in main execution: {str(e)}")
        return False
    finally:
        # Clean up
        spark.stop()


if __name__ == "__main__":
    main()
